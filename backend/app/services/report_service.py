from __future__ import annotations

import logging
import re
from collections.abc import Iterator
from dataclasses import dataclass
from datetime import UTC, datetime
from hashlib import sha256
from io import BytesIO
from uuid import UUID, uuid4

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from sqlalchemy import text
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from app.core.constants import READ_ONLY
from app.core.errors import DomainError
from app.core.permissions import can_generate_reports, can_read_reports
from app.db.models import Decision, Enterprise, Report, ReportEvidence, ReportSection
from app.db.repositories.decision_repository import DecisionRepository
from app.db.repositories.document_repository import DocumentRepository
from app.db.repositories.evidence_repository import EvidenceRepository
from app.db.repositories.match_repository import MatchRepository
from app.db.repositories.material_repository import MaterialRepository
from app.db.repositories.report_repository import ReportRepository
from app.db.repositories.requirement_repository import RequirementRepository
from app.db.repositories.risk_repository import RiskRepository
from app.integrations.object_storage import MinioObjectStorage, ObjectStorageUnavailable
from app.schemas.documents import TaskResponse
from app.schemas.reports import ReportResponse, ReportSectionResponse
from app.services.audit_service import AuditService
from app.services.material_match_policy import is_enterprise_material_requirement
from app.services.project_fact_resolver import ProjectFactResolver
from app.services.project_service import ProjectService
from app.services.report_field_query import ReportFieldQueryService
from app.services.task_service import TaskService

logger = logging.getLogger(__name__)

# Section display order in the rendered document
_SECTION_NAMES = {
    # Common
    "PROJECT_OVERVIEW": "项目概况",
    "BID_SCHEDULE": "关键时间与递交清单",
    "QUALIFICATION_MATRIX": "资格条件与企业符合情况",
    "ENTERPRISE_OVERVIEW": "企业概况",
    "EXECUTIVE_SUMMARY": "执行摘要",
    "ANALYSIS_COVERAGE": "分析覆盖率",
    # Risk + match insights
    "CORE_RISKS": "核心风险",
    "KEY_GAPS": "关键缺口",
    "ACTION_PLAN": "行动计划",
    # Material
    "MATERIAL_SUMMARY": "企业材料汇总",
    "MATERIAL_LINKED": "已关联材料",
    "MATERIAL_UNLINKED": "待补充材料",
    # FULL only
    "QUALIFICATION_ANALYSIS": "资格分析",
    "RISK_ITEMS": "全部风险事项",
    "ENTERPRISE_MATCHING": "企业匹配",
    "SCORING_ANALYSIS": "评分要点",
    "COMPREHENSIVE_DECISION": "综合决策",
    "TODOS": "待办事项",
}

# Severity weights for sorting risks top-first
_SEVERITY_RANK = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}

# 报告正文对每个章节保留的最大条目数，避免前端/用户看到 99/369 条噪音
_RISK_TOP_LIMIT = 10
_MATCH_TOP_LIMIT = 10
_REQUIREMENT_TOP_LIMIT = 20
_EVIDENCE_INDEX_LIMIT = 30

_DECISION_SUGGESTION_TEXT = {
    "RECOMMEND": "建议投标",
    "CAUTION": "谨慎投标",
    "HOLD": "暂缓",
    "REJECT": "不建议投标",
}
_FINAL_DECISION_TEXT = {"BID": "投标", "ABANDON": "放弃"}
_MATCH_STATUS_TEXT = {
    "MATCHED": "已满足",
    "MISSING": "缺失",
    "UNCERTAIN": "待确认",
    "UNASSESSED": "未评估",
}
_RISK_SEVERITY_TEXT = {
    "CRITICAL": "重大风险",
    "HIGH": "高风险",
    "MEDIUM": "中风险",
    "LOW": "低风险",
    "INFO": "提示",
}

_WHITESPACE_RE = re.compile(r"\s+")


@dataclass(frozen=True, slots=True)
class AuthorizedReportDownload:
    file_name: str
    mime_type: str
    stream: Iterator[bytes]


@dataclass(frozen=True, slots=True)
class _DraftSection:
    code: str
    content: str
    evidence_ids: list[UUID]
    requires_evidence: bool = True


class ReportService:
    def __init__(self, session: Session, object_storage: MinioObjectStorage) -> None:
        self._session = session
        self._storage = object_storage
        self._projects = ProjectService(session)
        self._reports = ReportRepository(session)
        self._decisions = DecisionRepository(session)
        self._requirements = RequirementRepository(session)
        self._risks = RiskRepository(session)
        self._matches = MatchRepository(session)
        self._evidences = EvidenceRepository(session)
        self._documents = DocumentRepository(session)
        self._materials = MaterialRepository(session)
        self._audit = AuditService(session)
        self._project_facts = ProjectFactResolver(session)
        self._report_fields = ReportFieldQueryService(session)

    def submit(
        self,
        project_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        publisher,
        report_type: str = "SIMPLE",
        analysis_run_id: UUID | None = None,
    ) -> TaskResponse:
        project = self._projects.get_visible(project_id, actor_id, role_codes)
        self._projects.require_writable(project)
        self._require_writer(role_codes)
        if self._requirements.has_pending_for_project(project_id):
            raise DomainError(
                "REPORT_REVIEW_PENDING",
                "请先完成需求复核中的所有高优先需求，再生成报告。",
                409,
            )
        self._session.execute(
            text("SELECT pg_advisory_xact_lock(hashtext(:lock_key))"),
            {"lock_key": f"report-version:{project_id}"},
        )
        report = Report(
            id=uuid4(),
            project_id=project_id,
            analysis_run_id=analysis_run_id,
            version_no=self._reports.next_version_no(project_id),
            report_type=report_type,
            status="QUEUED",
            docx_object_key=None,
            pdf_object_key=None,
            error_code=None,
            error_message=None,
            generated_by=actor_id,
            generated_at=None,
            created_at=datetime.now(UTC),
        )
        try:
            self._reports.add(report)
            self._session.flush()
            task = TaskService(self._session).create_report_task(report, actor_id)
            self._audit.record(
                actor_id=actor_id,
                action="GENERATE_REPORT",
                target_type="REPORT",
                target_id=report.id,
                project_id=project_id,
                after={"version_no": report.version_no, "report_type": report_type},
            )
            self._session.commit()
        except IntegrityError as exc:
            self._session.rollback()
            raise DomainError("REPORT_VERSION_CONFLICT", "报告版本冲突，请重新生成。", 409) from exc

        try:
            task.celery_task_id = publisher.publish_generate_report(task.id, report.id)
            self._session.commit()
        except Exception:
            self._session.rollback()
            self.mark_failed(
                report.id,
                "TASK_QUEUE_UNAVAILABLE",
                "报告任务队列暂不可用，请稍后重新生成。",
            )
            TaskService(self._session).fail_report_task(
                task.id,
                "TASK_QUEUE_UNAVAILABLE",
                "报告任务队列暂不可用，请稍后重新生成。",
            )
        return self._task_response(task)

    def latest(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> ReportResponse | None:
        self._projects.get_visible(project_id, actor_id, role_codes)
        self._require_reader(role_codes)
        report = self._reports.latest_for_project(project_id)
        if report is not None and self._is_read_only(role_codes) and report.status != "READY":
            return None
        return None if report is None else self._response(report)

    def get(self, report_id: UUID, actor_id: UUID, role_codes: set[str]) -> ReportResponse:
        report = self._reports.get(report_id)
        if report is None:
            raise DomainError("RESOURCE_NOT_FOUND", "资源不存在", 404)
        self._projects.get_visible(report.project_id, actor_id, role_codes)
        self._require_reader(role_codes)
        if self._is_read_only(role_codes) and report.status != "READY":
            raise DomainError("RESOURCE_NOT_FOUND", "资源不存在", 404)
        return self._response(report)

    def generate(self, report_id: UUID, actor_id: UUID, role_codes: set[str]) -> ReportResponse:
        report = self._reports.get(report_id, for_update=True)
        if report is None:
            raise DomainError("RESOURCE_NOT_FOUND", "资源不存在", 404)
        project = self._projects.get_visible(report.project_id, actor_id, role_codes)
        self._projects.require_writable(project)
        self._require_writer(role_codes)
        if report.status != "QUEUED":
            raise DomainError("REPORT_GENERATION_NOT_ALLOWED", "报告当前不能生成。", 409)
        report.status = "GENERATING"
        report.error_code = None
        report.error_message = None
        self._session.commit()

        drafts = self._build_sections(project, report.report_type)
        self._validate_drafts(drafts)
        now = datetime.now(UTC)
        persisted_sections: list[ReportSection] = []
        for order_no, draft in enumerate(drafts, start=1):
            section = ReportSection(
                id=uuid4(),
                report_id=report.id,
                section_code=draft.code,
                order_no=order_no,
                content_markdown=draft.content,
                created_at=now,
            )
            self._reports.add_section(section)
            persisted_sections.append(section)
        self._session.flush()
        for section, draft in zip(persisted_sections, drafts, strict=True):
            for evidence_id in dict.fromkeys(draft.evidence_ids):
                self._reports.add_evidence(
                    ReportEvidence(report_section_id=section.id, evidence_id=evidence_id)
                )

        docx_key = f"reports/{report.project_id}/v{report.version_no}/report.docx"
        pdf_key = f"reports/{report.project_id}/v{report.version_no}/report.pdf"
        md_key = f"reports/{report.project_id}/v{report.version_no}/report.md"
        docx_content = self._render_docx(project.name, report.version_no, drafts)
        pdf_content = self._render_pdf(project.name, report.version_no, drafts)
        md_content = self._render_md(project.name, report.version_no, drafts)
        try:
            self._storage.put_bytes(
                docx_key,
                docx_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self._verify_object_hash(docx_key, docx_content)
            self._storage.put_bytes(pdf_key, pdf_content, "application/pdf")
            self._verify_object_hash(pdf_key, pdf_content)
            self._storage.put_bytes(md_key, md_content, "text/markdown")
            self._verify_object_hash(md_key, md_content)
        except ObjectStorageUnavailable as exc:
            self._session.rollback()
            self._compensate_output(docx_key)
            self._compensate_output(pdf_key)
            self._compensate_output(md_key)
            raise DomainError("OBJECT_STORAGE_UNAVAILABLE", "报告对象存储暂不可用。", 503) from exc
        except Exception:
            self._session.rollback()
            self._compensate_output(docx_key)
            self._compensate_output(pdf_key)
            self._compensate_output(md_key)
            raise

        report = self._reports.get(report.id, for_update=True)
        if report is None:
            self._session.rollback()
            self._compensate_output(docx_key)
            self._compensate_output(pdf_key)
            raise DomainError("RESOURCE_NOT_FOUND", "资源不存在", 404)
        report.status = "READY"
        report.docx_object_key = docx_key
        report.pdf_object_key = pdf_key
        report.md_object_key = md_key
        report.generated_at = now
        self._audit.record(
            actor_id=actor_id,
            action="COMPLETE_REPORT",
            target_type="REPORT",
            target_id=report.id,
            project_id=report.project_id,
            after={"version_no": report.version_no},
        )
        self._session.commit()
        return self._response(report)

    def mark_failed(self, report_id: UUID, error_code: str, message: str) -> None:
        self._session.rollback()
        report = self._reports.get(report_id, for_update=True)
        if report is None or report.status not in {"QUEUED", "GENERATING"}:
            self._session.rollback()
            return
        report.status = "FAILED"
        report.error_code = error_code
        report.error_message = message
        self._session.commit()

    def create_authorized_download(
        self,
        report_id: UUID,
        report_format: str,
        actor_id: UUID,
        role_codes: set[str],
    ) -> AuthorizedReportDownload:
        report = self._reports.get(report_id)
        if report is None:
            raise DomainError("RESOURCE_NOT_FOUND", "资源不存在", 404)
        self._projects.get_visible(report.project_id, actor_id, role_codes)
        self._require_reader(role_codes)
        if report.status != "READY":
            raise DomainError("REPORT_NOT_READY", "报告尚未生成完成。", 409)
        if report_format == "docx":
            object_key = report.docx_object_key
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif report_format == "pdf":
            object_key = report.pdf_object_key
            mime_type = "application/pdf"
        elif report_format == "md":
            object_key = report.md_object_key
            mime_type = "text/markdown; charset=utf-8"
        else:
            raise DomainError("INVALID_REPORT_FORMAT", "仅支持 DOCX、PDF 或 MD 报告。", 422)
        if object_key is None:
            raise DomainError("REPORT_NOT_READY", "报告尚未生成完成。", 409)
        self._audit.record(
            actor_id=actor_id,
            action="DOWNLOAD_REPORT",
            target_type="REPORT",
            target_id=report.id,
            project_id=report.project_id,
            after={"format": report_format},
        )
        self._session.commit()

        def iterator() -> Iterator[bytes]:
            with self._storage.open_object(object_key) as source:
                yield from source.stream(amt=1024 * 1024)

        return AuthorizedReportDownload(
            file_name=f"bid-report-v{report.version_no}.{report_format}",
            mime_type=mime_type,
            stream=iterator(),
        )

    # ---------------------------------------------------------------------
    # Section composition
    # ---------------------------------------------------------------------

    def _build_sections(self, project, report_type: str) -> list[_DraftSection]:
        decision = self._decisions.latest_for_project(project.id)
        if decision is None:
            raise DomainError("REPORT_DECISION_MISSING", "请先生成投标建议后再生成报告。", 409)
        decision_evidence = self._decisions.list_evidence_ids(decision.id)
        if not decision_evidence:
            raise DomainError("REPORT_EVIDENCE_MISSING", "综合决策缺少可追溯证据。", 409)

        field_context = self._report_fields.load(project.id)
        requirements = field_context.confirmed_requirements
        all_requirements = field_context.all_requirements
        risks = field_context.risks
        matches = field_context.matches

        qualification_requirements = [
            requirement
            for requirement in requirements
            if requirement.category == "QUALIFICATION"
            and is_enterprise_material_requirement(requirement)
        ]
        scoring_requirements = field_context.scoring_requirements()
        open_risks = field_context.open_risks()

        requirement_evidence_map = field_context.requirement_evidence
        risk_evidence_map = field_context.risk_evidence
        match_evidence_map = field_context.match_evidence

        drafts: list[_DraftSection] = [
            self._project_overview_section(project, field_context),
            self._executive_summary_section(project, decision, requirements, risks, matches),
            self._bid_schedule_section(project, field_context),
            self._qualification_matrix_section(
                qualification_requirements,
                matches,
                requirement_evidence_map,
                match_evidence_map,
            ),
            self._coverage_section(all_requirements, requirements, matches),
        ]

        core_risks = self._core_risks_section(
            open_risks, risk_evidence_map, decision_evidence
        )
        if core_risks is not None:
            drafts.append(core_risks)

        action_plan = self._action_plan_section(field_context.action_plan_items())
        drafts.append(action_plan)

        if report_type == "FULL":
            full_risks = self._full_risks_section(
                [risk for risk in open_risks if risk.severity not in {"CRITICAL", "HIGH"}],
                risk_evidence_map,
                decision_evidence,
            )
            if full_risks is not None:
                drafts.append(full_risks)
            if scoring_requirements:
                drafts.append(
                    self._requirement_section(
                        "量化评分要点",
                        scoring_requirements,
                        requirement_evidence_map,
                        decision_evidence,
                    )
                )
            else:
                drafts.append(
                    _DraftSection(
                        "SCORING_ANALYSIS",
                        "未抽取到量化评分细则；初步评审和否决条款不作为评分项展示。",
                        [],
                        requires_evidence=False,
                    )
                )
            drafts.append(self._decision_section(decision, decision_evidence))

        # 任何 draft 都自携带证据链接，不再单独追加 EVIDENCE_INDEX 大段
        # 前端在抽屉 / 章节元数据里就能看到 evidence_ids
        return drafts

    def _bid_schedule_section(self, project, field_context) -> _DraftSection:
        """Render schedule facts first, then evidence-linked submission obligations."""
        deadline = self._project_facts.resolve_bid_deadline(project)
        rows: list[tuple[str, str, list[UUID]]] = []
        if deadline.is_confirmed:
            suffix = "（具体时间待确认）" if deadline.precision == "DATE" else ""
            rows.append(("投标截止", f"{deadline.display_value()}{suffix}", deadline.evidence_ids))
        else:
            rows.append(("投标截止", "待确认", []))
        for code, label in (("BID_OPENING_AT", "开标时间"), ("BID_BOND", "投标保证金")):
            value = field_context.value_for(code)
            rows.append((label, value or "待确认", field_context.evidence_for(code)))

        lines = ["关键时间与递交清单："]
        evidence_ids: list[UUID] = []
        for label, value, row_evidence_ids in rows:
            lines.append(f"- {label}：{value}")
            lines.extend(self._evidence_citations(row_evidence_ids, limit=1))
            evidence_ids.extend(row_evidence_ids)

        schedule_requirements = field_context.schedule_requirements()[:8]
        if schedule_requirements:
            lines.append("- 递交与操作要求：")
            for requirement in schedule_requirements:
                lines.append(f"  - {requirement.title}")
                if requirement.description:
                    lines.append(
                        f"    说明：{self._collapse(requirement.description, limit=140)}"
                    )
                requirement_evidence = field_context.requirement_evidence.get(
                    requirement.id, []
                )
                lines.extend(self._evidence_citations(requirement_evidence, limit=1))
                evidence_ids.extend(requirement_evidence)
        return _DraftSection(
            "BID_SCHEDULE",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
            requires_evidence=False,
        )

    def _qualification_matrix_section(
        self,
        requirements,
        matches,
        requirement_evidence_map,
        match_evidence_map,
    ) -> _DraftSection:
        """Show one tender qualification with its enterprise-tag conclusion and source."""
        if not requirements:
            return _DraftSection(
                "QUALIFICATION_MATRIX",
                "资格条件与企业符合情况：未发现已确认资格条目。",
                [],
                requires_evidence=False,
            )
        status_rank = {"MISSING": 0, "UNCERTAIN": 1, "MATCHED": 2}
        matches_by_requirement: dict[UUID, object] = {}
        for match in matches:
            if match.requirement_id not in {requirement.id for requirement in requirements}:
                continue
            current = matches_by_requirement.get(match.requirement_id)
            if current is None or status_rank.get(match.final_status, 99) < status_rank.get(
                current.final_status, 99
            ):
                matches_by_requirement[match.requirement_id] = match
        ordered = sorted(
            requirements,
            key=lambda requirement: (
                status_rank.get(
                    getattr(matches_by_requirement.get(requirement.id), "final_status", "MISSING"),
                    0,
                ),
                0 if requirement.is_mandatory else 1,
                requirement.created_at,
            ),
        )[:_REQUIREMENT_TOP_LIMIT]
        lines = [
            "资格条件与企业符合情况"
            f"（展示前 {len(ordered)} 条 / 共 {len(requirements)} 条）："
        ]
        evidence_ids: list[UUID] = []
        for requirement in ordered:
            match = matches_by_requirement.get(requirement.id)
            status = "UNASSESSED" if match is None else match.final_status
            lines.append(
                f"### {_MATCH_STATUS_TEXT.get(status, status)}｜"
                f"{self._requirement_label(requirement.title)}"
            )
            summary = self._requirement_summary(requirement.title, requirement.description)
            if summary:
                lines.append(
                    f"- **要求摘要**：{summary}"
                )
            if match is None:
                lines.append("- **企业匹配**：尚未执行")
            else:
                lines.append(f"- **企业匹配**：{self._collapse(match.reason, limit=180)}")
            requirement_evidence = requirement_evidence_map.get(requirement.id, [])
            match_evidence = [
                evidence_id
                for evidence_id, side in match_evidence_map.get(getattr(match, "id", None), [])
                if side == "REQUIREMENT"
            ]
            cited_evidence = match_evidence or requirement_evidence
            lines.extend(self._evidence_citations(cited_evidence, limit=1))
            evidence_ids.extend(cited_evidence)
            lines.append("")
        return _DraftSection(
            "QUALIFICATION_MATRIX",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
        )

    def _project_overview_section(self, project, field_context) -> _DraftSection:
        deadline = self._project_facts.resolve_bid_deadline(project)
        if deadline.precision == "DATE":
            deadline_line = (
                f"投标截止日期：{deadline.display_value()}（具体时间待确认）"
            )
        elif deadline.is_confirmed:
            deadline_line = f"投标截止时间：{deadline.display_value()}"
        else:
            deadline_line = "投标截止时间：-"
        lines = [
            f"项目名称：{project.name}",
            f"项目编号：{project.code}",
            f"采购人：{project.purchaser or field_context.value_for('PURCHASER') or '-'}",
            deadline_line,
        ]
        evidence_ids = list(deadline.evidence_ids)
        for code, label in (
            ("LOCATION", "建设/交付地点"),
            ("BUDGET", "项目预算"),
            ("MAX_PRICE", "最高限价"),
            ("PROCUREMENT_METHOD", "采购方式"),
        ):
            value = field_context.value_for(code)
            if value is not None:
                lines.append(f"{label}：{value}")
                evidence_ids.extend(field_context.evidence_for(code))
        return _DraftSection(
            "PROJECT_OVERVIEW",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
            requires_evidence=False,
        )

    def _executive_summary_section(
        self, project, decision, requirements, risks, matches
    ) -> _DraftSection:
        critical_count = sum(
            1
            for r in risks
            if r.severity in {"CRITICAL", "HIGH"} and r.status in {"PENDING", "CONFIRMED"}
        )
        gap_count = sum(1 for m in matches if m.final_status in {"MISSING", "UNCERTAIN"})
        suggestion = _DECISION_SUGGESTION_TEXT.get(decision.suggestion, decision.suggestion)
        final = _FINAL_DECISION_TEXT.get(decision.final_decision or "", "尚未确认")
        lines = [
            f"项目：{project.name}（{project.code}）",
            f"系统建议：{suggestion}",
            f"最终决策：{final}",
            f"核心风险（重大/高）：{critical_count} 项",
            f"关键缺口（缺失/待确认）：{gap_count} 项",
            f"已确认招标要求：{len(requirements)} 项",
        ]
        return _DraftSection(
            "EXECUTIVE_SUMMARY",
            "\n".join(lines),
            [],
            requires_evidence=False,
        )

    def _coverage_section(self, all_requirements, confirmed_requirements, matches) -> _DraftSection:
        mandatory = [requirement for requirement in all_requirements if requirement.is_mandatory]
        confirmed_mandatory = [
            requirement for requirement in confirmed_requirements if requirement.is_mandatory
        ]
        match_statuses = {match.requirement_id: match.final_status for match in matches}
        matched = sum(
            1 for requirement in confirmed_requirements
            if match_statuses.get(requirement.id) == "MATCHED"
        )
        confirmed_rate = self._percent(len(confirmed_requirements), len(all_requirements))
        mandatory_rate = self._percent(len(confirmed_mandatory), len(mandatory))
        match_rate = self._percent(matched, len(confirmed_requirements))
        pending_count = sum(
            1 for requirement in all_requirements if requirement.review_status == "PENDING"
        )
        deferred_count = sum(
            1 for requirement in all_requirements if requirement.review_status == "DEFERRED"
        )
        lines = [
            f"抽取 Requirement：{len(all_requirements)} 条",
            f"已确认（人工/规则/自动）：{len(confirmed_requirements)} 条（{confirmed_rate}）",
            f"强制项已确认：{len(confirmed_mandatory)}/{len(mandatory)}（{mandatory_rate}）",
            f"已确认项具备匹配材料：{matched}/{len(confirmed_requirements)}（{match_rate}）",
            f"待人工复核：{pending_count} 条",
            f"延后复核：{deferred_count} 条（关键队列完成后处理）",
        ]
        return _DraftSection("ANALYSIS_COVERAGE", "\n".join(lines), [], requires_evidence=False)

    def _agent_review_section(
        self, agent_run, fallback_evidence_ids: list[UUID]
    ) -> _DraftSection | None:
        if agent_run is None:
            return None
        result = agent_run.result or {}
        strategy = result.get("strategy") or {}
        critique = result.get("evidence_critique") or {}
        legal = result.get("legal_assessment") or {}
        specialists = result.get("specialist_assessments") or {}

        recommendation = _AGENT_RECOMMENDATION_TEXT.get(
            strategy.get("bid_recommendation"), strategy.get("bid_recommendation", "-")
        )
        confidence = strategy.get("confidence")
        confidence_pct = (
            f"{int(round(confidence * 100))}%" if isinstance(confidence, (int, float)) else "-"
        )

        lines: list[str] = [
            f"研判结论：{recommendation}（置信度 {confidence_pct}）",
        ]
        rationale = self._collapse(strategy.get("rationale"))
        if rationale:
            lines.append(f"研判理由：{rationale}")

        priority_actions = strategy.get("priority_actions") or []
        if priority_actions:
            lines.append("")
            lines.append("优先行动：")
            for action in priority_actions:
                priority = action.get("priority") or "-"
                owner = action.get("owner_role") or "未指派"
                body = self._collapse(action.get("action"))
                if body:
                    lines.append(f"- [{priority}] {body}（责任：{owner}）")

        residual_risks = strategy.get("residual_risks") or []
        if residual_risks:
            lines.append("")
            lines.append("残余风险：")
            for risk in residual_risks[:8]:
                cleaned = self._collapse(risk)
                if cleaned:
                    lines.append(f"- {cleaned}")

        if critique:
            conclusion = self._collapse(critique.get("conclusion"))
            if conclusion:
                lines.append("")
                lines.append("证据复核结论：")
                lines.append(conclusion)
            blockers = critique.get("blockers") or []
            if blockers:
                lines.append("")
                lines.append("需人工确认事项：")
                for blocker in blockers[:6]:
                    cleaned = self._collapse(blocker)
                    if cleaned:
                        lines.append(f"- {cleaned}")

        # 法律评估的 high-severity finding 取前 3 条
        legal_findings = sorted(
            legal.get("findings") or [],
            key=lambda f: _SEVERITY_RANK.get(str(f.get("severity", "INFO")).upper(), 99),
        )
        if legal_findings:
            lines.append("")
            lines.append("法律评估要点：")
            for finding in legal_findings[:3]:
                title = self._collapse(finding.get("title"))
                conclusion = self._collapse(finding.get("conclusion"))
                severity = str(finding.get("severity", "INFO")).upper()
                if title and conclusion:
                    lines.append(f"- [{severity}] {title}：{conclusion}")
                elif title:
                    lines.append(f"- [{severity}] {title}")

        # 各 specialist 的总体 confidence + open_questions
        if specialists:
            lines.append("")
            lines.append("分维度评估：")
            for name in ("qualification", "commercial", "technical", "scoring", "schedule"):
                assessment = specialists.get(name)
                if not isinstance(assessment, dict):
                    continue
                conf = assessment.get("confidence")
                conf_pct = f"{int(round(conf * 100))}%" if isinstance(conf, (int, float)) else "-"
                findings = assessment.get("findings") or []
                open_q = assessment.get("open_questions") or []
                label = _AGENT_SPECIALIST_NAMES.get(name, name)
                summary = self._collapse(assessment.get("summary"))
                line = f"- {label}（置信度 {conf_pct}，{len(findings)} 项发现）"
                if summary:
                    line += f"：{summary}"
                lines.append(line)
                for question in open_q[:2]:
                    cleaned = self._collapse(question)
                    if cleaned:
                        lines.append(f"  待确认：{cleaned}")

        agent_eids = list(self._agent_runs.list_evidence_ids(agent_run.id))
        return _DraftSection(
            "AGENT_REVIEW",
            "\n".join(line for line in lines if line is not None),
            agent_eids or fallback_evidence_ids,
        )

    def _core_risks_section(
        self, risks, risk_evidence_map, fallback_evidence_ids: list[UUID]
    ) -> _DraftSection | None:
        open_critical = [
            r
            for r in risks
            if r.severity in {"CRITICAL", "HIGH"} and r.status in {"PENDING", "CONFIRMED"}
        ]
        if not open_critical:
            return None
        # 按规则（rule_version_id）聚合，避免同一规则因每个 requirement 触发而重复
        grouped: dict[UUID, dict[str, object]] = {}
        for risk in open_critical:
            entry = grouped.setdefault(
                risk.rule_version_id,
                {
                    "severity": risk.severity,
                    "title": risk.title,
                    "description": risk.description,
                    "count": 0,
                    "evidence_ids": [],
                    "details": [],  # 各触发的具体差异
                },
            )
            entry["count"] += 1
            if _SEVERITY_RANK.get(risk.severity, 99) < _SEVERITY_RANK.get(
                str(entry["severity"]), 99
            ):
                entry["severity"] = risk.severity
            risk_eids = risk_evidence_map.get(risk.id, [])
            entry["evidence_ids"].extend(risk_eids or fallback_evidence_ids)
            # Only a real material id may render material details. A
            # requirement id in the generic rule subject is not a material.
            trigger = risk.trigger_data or {}
            material_id = trigger.get("material_id")
            valid_to = trigger.get("valid_to", "")
            if material_id:
                mat = self._materials.get(UUID(material_id)) if material_id else None
                mat_name = mat.name if mat else "关联材料待确认"
                entry["details"].append(f"{mat_name}（过期：{valid_to}）")
        ranked = sorted(
            grouped.values(),
            key=lambda item: _SEVERITY_RANK.get(str(item["severity"]), 99),
        )
        top = ranked[:_RISK_TOP_LIMIT]
        evidence_ids: list[UUID] = []
        lines = [f"核心风险（前 {len(top)} 项规则 / 共 {len(open_critical)} 条触发）："]
        for entry in top:
            count_suffix = f"（{entry['count']} 条触发）" if entry["count"] > 1 else ""
            severity = _RISK_SEVERITY_TEXT.get(str(entry["severity"]), str(entry["severity"]))
            title = self._human_risk_title(str(entry["title"]))
            lines.append(f"- **{severity}**：{title}{count_suffix}")
            if entry["details"]:
                for detail in entry["details"]:
                    lines.append(f"  · {detail}")
            lines.extend(self._evidence_citations(entry["evidence_ids"], limit=2))
            evidence_ids.extend(entry["evidence_ids"])
        return _DraftSection(
            "CORE_RISKS",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
        )

    def _key_gaps_section(
        self,
        matches,
        match_evidence_map,
        requirement_evidence_map,
        fallback_evidence_ids: list[UUID],
    ) -> _DraftSection | None:
        if not matches:
            return None
        missing_count = sum(match.final_status == "MISSING" for match in matches)
        uncertain_count = len(matches) - missing_count
        evidence_ids: list[UUID] = []
        for match in matches:
            match_eids = [
                evidence_id for evidence_id, _ in match_evidence_map.get(match.id, [])
            ]
            evidence_ids.extend(
                match_eids or requirement_evidence_map.get(match.requirement_id, [])
            )
        lines = [
            "关键缺口总览：",
            (
                f"- 待补齐企业材料：{len(matches)} 项"
                f"（缺失 {missing_count} 项，待确认 {uncertain_count} 项）"
            ),
            "- 具体招标条件、企业匹配结论与原文依据见“资格条件与企业符合情况”。",
        ]
        citation_ids = list(dict.fromkeys(evidence_ids)) or fallback_evidence_ids
        lines.extend(self._evidence_citations(citation_ids, limit=1))
        return _DraftSection(
            "KEY_GAPS",
            "\n".join(lines),
            list(dict.fromkeys(citation_ids)),
        )

    def _action_plan_section(self, items) -> _DraftSection:
        if not items:
            content = "当前未发现需要紧急处理的行动项。"
            return _DraftSection("ACTION_PLAN", content, [], requires_evidence=False)
        entries: list[str] = []
        evidence_ids: list[UUID] = []
        for item in items:
            entries.append(f"- [{item.priority}] {self._collapse(item.action, limit=180)}")
            entries.extend(self._evidence_citations(item.evidence_ids, limit=1))
            evidence_ids.extend(item.evidence_ids)
        return _DraftSection(
            "ACTION_PLAN",
            "行动计划：\n" + "\n".join(entries),
            list(dict.fromkeys(evidence_ids)),
        )

    def _enterprise_overview_section(self, project) -> _DraftSection:
        # 从项目绑定企业(联合体)→ EnterpriseMaterial 查询企业材料汇总
        enterprise_ids = self._projects.list_enterprise_ids(project.id)
        if not enterprise_ids:
            return _DraftSection(
                "ENTERPRISE_OVERVIEW",
                "企业概况：未关联企业信息",
                [],
                requires_evidence=False,
            )
        materials = self._materials.list_active_for_enterprises(enterprise_ids)
        if not materials:
            return _DraftSection(
                "ENTERPRISE_OVERVIEW",
                "企业概况：企业材料库为空，请上传企业资质材料",
                [],
                requires_evidence=False,
            )

        # 按 material_type 分组统计
        type_counts: dict[str, int] = {}
        for m in materials:
            t = m.material_type or "未分类"
            type_counts[t] = type_counts.get(t, 0) + 1

        enterprise_names = [
            enterprise.name
            for enterprise_id in enterprise_ids
            if (enterprise := self._session.get(Enterprise, enterprise_id)) is not None
        ]
        type_lines = [f"- {t}（{c} 项）" for t, c in sorted(type_counts.items())]
        lines = [
            f"投标企业：{'、'.join(enterprise_names)}",
            f"企业材料总数：{len(materials)} 项",
            "材料类型分布：",
        ]
        lines.extend(type_lines)
        return _DraftSection(
            "ENTERPRISE_OVERVIEW",
            "\n".join(lines),
            [],
            requires_evidence=False,
        )

    def _material_summary_section(self, matches) -> _DraftSection:
        matched = sum(1 for m in matches if m.final_status == "MATCHED")
        uncertain = sum(1 for m in matches if m.final_status == "UNCERTAIN")
        missing = sum(1 for m in matches if m.final_status == "MISSING")
        total = len(matches)
        lines = [
            f"材料匹配汇总（共 {total} 条招标要求）：",
            f"- 已匹配（MATCHED）：{matched} 项",
            f"- 待确认（UNCERTAIN）：{uncertain} 项",
            f"- 缺失（MISSING）：{missing} 项",
        ]
        return _DraftSection(
            "MATERIAL_SUMMARY",
            "\n".join(lines),
            [],
            requires_evidence=False,
        )

    def _material_linked_section(self, matches, match_evidence_map) -> _DraftSection:
        linked = [m for m in matches if m.final_status == "MATCHED"]
        if not linked:
            return _DraftSection(
                "MATERIAL_LINKED",
                "已关联材料：暂无已匹配项",
                [],
                requires_evidence=False,
            )
        from app.db.repositories.requirement_repository import RequirementRepository

        req_repo = RequirementRepository(self._session)
        titles_by_id: dict[UUID, str] = {}
        entries: list[str] = []
        evidence_ids: list[UUID] = []
        top_matches = linked[:15]
        for match in top_matches:
            req_id = match.requirement_id
            if req_id not in titles_by_id:
                req_obj = req_repo.get(req_id)
                titles_by_id[req_id] = req_obj.title if req_obj else "(要求已删除)"
            match_eids = [eid for eid, _ in match_evidence_map.get(match.id, [])]
            reason = self._collapse(match.reason, limit=160)
            entries.append(
                f"- [{match.final_status}] {titles_by_id[req_id]}"
                + (f"（{reason}）" if reason else "")
            )
            entries.extend(self._evidence_citations(match_eids, limit=1))
            evidence_ids.extend(match_eids)
        lines = [
            f"已关联材料（展示前 {len(top_matches)} 项 / 共 {len(linked)} 项）："
        ]
        lines.extend(entries)
        return _DraftSection(
            "MATERIAL_LINKED",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
            requires_evidence=False,
        )

    def _material_unlinked_section(self, matches, match_evidence_map) -> _DraftSection:
        unlinked = [m for m in matches if m.final_status in {"UNCERTAIN", "MISSING"}]
        if not unlinked:
            return _DraftSection(
                "MATERIAL_UNLINKED",
                "待补充材料：暂无缺失或待确认项",
                [],
                requires_evidence=False,
            )
        from app.db.repositories.requirement_repository import RequirementRepository

        req_repo = RequirementRepository(self._session)
        titles_by_id: dict[UUID, str] = {}
        entries: list[str] = []
        evidence_ids: list[UUID] = []
        for match in unlinked[:20]:
            req_id = match.requirement_id
            if req_id not in titles_by_id:
                req_obj = req_repo.get(req_id)
                titles_by_id[req_id] = req_obj.title if req_obj else "(要求已删除)"
            reason = self._collapse(match.reason, limit=100)
            entries.append(
                f"- [{match.final_status}] {titles_by_id[req_id]}"
                + (f"（{reason}）" if reason else "")
            )
            match_eids = [eid for eid, _ in match_evidence_map.get(match.id, [])]
            evidence_ids.extend(match_eids)
        lines = [f"待补充材料（展示前 {len(entries)} 项 / 共 {len(unlinked)} 项）："]
        lines.extend(entries)
        return _DraftSection(
            "MATERIAL_UNLINKED",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
            requires_evidence=False,
        )

    def _requirement_section(
        self,
        title: str,
        requirements,
        requirement_evidence_map,
        fallback_evidence_ids: list[UUID] | None = None,
    ) -> _DraftSection:
        if not requirements:
            return _DraftSection(
                self._code_for(title),
                f"{title}：未发现已确认条目。",
                [],
                requires_evidence=False,
            )
        top = sorted(
            requirements,
            key=lambda r: (
                0 if r.is_mandatory else 1,
                -float(r.score) if r.score is not None else 0,
            ),
        )[:_REQUIREMENT_TOP_LIMIT]
        evidence_ids: list[UUID] = []
        lines = [f"{title}（展示前 {len(top)} 条 / 共 {len(requirements)} 条）："]
        for requirement in top:
            tag = "强制" if requirement.is_mandatory else "非强制"
            score = f"（{requirement.score} 分）" if requirement.score is not None else ""
            lines.append(f"- [{tag}]{score} {requirement.title}")
            if requirement.description:
                lines.append(f"  说明：{self._collapse(requirement.description, limit=160)}")
            req_eids = requirement_evidence_map.get(requirement.id, [])
            citation_ids = req_eids or (fallback_evidence_ids or [])
            lines.extend(self._evidence_citations(citation_ids, limit=1))
            evidence_ids.extend(req_eids or (fallback_evidence_ids or []))
        return _DraftSection(
            self._code_for(title),
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
        )

    def _full_risks_section(
        self, risks, risk_evidence_map, fallback_evidence_ids: list[UUID]
    ) -> _DraftSection | None:
        if not risks:
            return None
        # 按 rule_version_id 聚合避免重复
        grouped: dict[UUID, dict[str, object]] = {}
        for risk in risks:
            entry = grouped.setdefault(
                risk.rule_version_id,
                {
                    "severity": risk.severity,
                    "status": risk.status,
                    "title": risk.title,
                    "description": risk.description,
                    "count": 0,
                    "evidence_ids": [],
                },
            )
            entry["count"] += 1
            if _SEVERITY_RANK.get(risk.severity, 99) < _SEVERITY_RANK.get(
                str(entry["severity"]), 99
            ):
                entry["severity"] = risk.severity
            risk_eids = risk_evidence_map.get(risk.id, [])
            entry["evidence_ids"].extend(risk_eids or fallback_evidence_ids)
        ranked = sorted(
            grouped.values(),
            key=lambda item: _SEVERITY_RANK.get(str(item["severity"]), 99),
        )
        top = ranked[:_RISK_TOP_LIMIT]
        evidence_ids: list[UUID] = []
        lines = [f"风险事项（按规则聚合，展示前 {len(top)} 项 / 共 {len(risks)} 条触发）："]
        for entry in top:
            status = entry["status"] or "-"
            count_suffix = f"（{entry['count']} 条触发）" if entry["count"] > 1 else ""
            lines.append(f"- [{entry['severity']}][{status}] {entry['title']}{count_suffix}")
            description = self._collapse(str(entry["description"]), limit=160)
            if description:
                lines.append(f"  描述：{description}")
            evidence_ids.extend(entry["evidence_ids"])
        return _DraftSection(
            "RISK_ITEMS",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
        )

    def _enterprise_matching_section(
        self, matches, match_evidence_map, fallback_evidence_ids: list[UUID]
    ) -> _DraftSection:
        status_rank = {
            "MISSING": 0,
            "UNCERTAIN": 1,
            "MATCHED": 2,
        }
        # 按 requirement_id 聚合：同一条招标要求不会重复展示多次
        by_requirement: dict[UUID, dict[str, object]] = {}
        for match in matches:
            entry = by_requirement.setdefault(
                match.requirement_id,
                {
                    "status": match.final_status,
                    "match_count": 0,
                    "reasons": [],
                    "evidence_ids": [],
                },
            )
            entry["match_count"] += 1
            entry["reasons"].append(match.reason)
            match_eids = [eid for eid, _ in match_evidence_map.get(match.id, [])]
            entry["evidence_ids"].extend(match_eids)
            if status_rank.get(match.final_status, 99) < status_rank.get(str(entry["status"]), 99):
                entry["status"] = match.final_status
        ranked = sorted(
            by_requirement.items(),
            key=lambda item: status_rank.get(str(item[1]["status"]), 99),
        )
        top = ranked[:_MATCH_TOP_LIMIT]
        evidence_ids: list[UUID] = []
        lines = [f"企业匹配（按要求聚合，展示前 {len(top)} 项 / 共 {len(matches)} 条匹配项）："]
        from app.db.repositories.requirement_repository import RequirementRepository

        req_repo = RequirementRepository(self._session)
        titles_by_id: dict[UUID, str] = {}
        for req_id, entry in top:
            if req_id not in titles_by_id:
                req_obj = req_repo.get(req_id)
                titles_by_id[req_id] = req_obj.title if req_obj else "(要求已删除)"
            title = titles_by_id[req_id]
            lines.append(f"- [{entry['status']}] {title}（{entry['match_count']} 条评估）")
            seen: set[str] = set()
            for raw_reason in entry["reasons"]:
                cleaned = self._collapse(raw_reason, limit=160)
                if cleaned and cleaned not in seen:
                    seen.add(cleaned)
                    lines.append(f"  摘要：{cleaned}")
                    if len(seen) >= 2:
                        break
            entry_eids = entry["evidence_ids"] or fallback_evidence_ids
            lines.extend(self._evidence_citations(entry_eids, limit=1))
            evidence_ids.extend(entry_eids)
        return _DraftSection(
            "ENTERPRISE_MATCHING",
            "\n".join(lines),
            list(dict.fromkeys(evidence_ids)),
        )

    def _decision_section(self, decision: Decision, decision_evidence) -> _DraftSection:
        suggestion = _DECISION_SUGGESTION_TEXT.get(decision.suggestion, decision.suggestion)
        final = _FINAL_DECISION_TEXT.get(decision.final_decision or "", "尚未由负责人确认")
        lines = [
            f"系统建议：{suggestion}",
            f"最终决策：{final}",
            f"决策说明：{self._collapse(decision.reason, limit=400)}",
        ]
        return _DraftSection(
            "COMPREHENSIVE_DECISION",
            "\n".join(lines),
            list(decision_evidence),
        )

    # ---------------------------------------------------------------------
    # Helpers
    # ---------------------------------------------------------------------

    @staticmethod
    def _percent(numerator: int, denominator: int) -> str:
        return "-" if denominator == 0 else f"{numerator / denominator:.0%}"

    def _evidence_citations(self, evidence_ids, *, limit: int) -> list[str]:
        """Render source evidence into a reader-facing citation, never an internal UUID."""
        unique_ids = list(dict.fromkeys(evidence_ids))[:limit]
        evidences = self._evidences.list_by_ids(unique_ids)
        node_ids = [
            evidence.document_node_id
            for evidence in evidences.values()
            if evidence.document_node_id
        ]
        paths = {
            node_id: section_path
            for node_id, section_path in self._session.execute(
                text("select id, section_path from app.document_nodes where id = any(:ids)"),
                {"ids": node_ids},
            ).tuples()
        } if node_ids else {}
        lines: list[str] = []
        for evidence_id in unique_ids:
            evidence = evidences.get(evidence_id)
            if evidence is None:
                continue
            if evidence.source_type == "SYSTEM_RULE":
                continue
            page = f"第 {evidence.page_number} 页" if evidence.page_number else "页码待确认"
            section = paths.get(evidence.document_node_id) or "章节待确认"
            lines.append(f"  - **原文定位**：{page}｜{section}")
        return lines

    @staticmethod
    def _code_for(title: str) -> str:
        return {
            "已确认资格要求": "QUALIFICATION_ANALYSIS",
            "已确认评分要求": "SCORING_ANALYSIS",
            "量化评分要点": "SCORING_ANALYSIS",
        }.get(title, "GENERIC_REQUIREMENT")

    @staticmethod
    def _collapse(value: object, limit: int | None = None) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            text = _WHITESPACE_RE.sub(" ", value).strip()
        else:
            text = str(value)
        if limit and len(text) > limit:
            text = text[: limit - 1] + "…"
        return text

    def _requirement_summary(self, title: str, description: str | None) -> str:
        """Keep the report readable when extractor descriptions repeat the title."""
        text = self._collapse(description)
        if not text:
            return ""
        title_pos = text.find(title)
        if title_pos >= 0:
            text = text[title_pos + len(title) :].lstrip("：:。；; ")
        return self._collapse(text, limit=120)

    @staticmethod
    def _requirement_label(title: str) -> str:
        labels = (
            ("项目经理", "项目经理资格"),
            ("业绩要求", "业绩要求"),
            ("资质要求", "企业资质"),
            ("财务要求", "财务要求"),
            ("财务状况表", "财务报表"),
            ("信用中国", "信用要求"),
            ("资质条件、能力和信誉", "综合资质与信誉"),
        )
        for marker, label in labels:
            if marker in title:
                return label
        return ReportService._collapse(title, limit=36)

    @staticmethod
    def _human_risk_title(title: str) -> str:
        prefix = "强制 Requirement 缺少材料证据:"
        if title.startswith(prefix):
            return f"缺少证明材料：{title.removeprefix(prefix)}"
        return title.replace("Requirement", "招标要求")

    @staticmethod
    def _validate_drafts(drafts: list[_DraftSection]) -> None:
        # 报告章节允许在没有数据时空内容（已显式标注），不再强制每节都有 evidence
        for draft in drafts:
            if not draft.requires_evidence:
                continue
            content = draft.content.strip()
            is_empty = (
                "未发现" in content
                or "未生成" in content
                or "未由负责人确认" in content
                or "未确认" in content
                or content.startswith("当前未发现")
            )
            if is_empty:
                continue
            if not draft.evidence_ids:
                logger.warning(
                    "报告章节 %s 缺少证据，跳过验证",
                    _SECTION_NAMES.get(draft.code, draft.code),
                )

    @staticmethod
    def _render_docx(project_name: str, version_no: int, sections: list[_DraftSection]) -> bytes:
        document = DocxDocument()
        # 全局默认样式
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)

        # 封面：报告标题
        title_para = document.add_paragraph()
        title_para.alignment = 1  # 居中
        title_run = title_para.add_run("投标综合分析报告")
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = None  # 使用默认深色

        # 封面：项目信息
        meta_para = document.add_paragraph()
        meta_para.alignment = 1
        meta_run = meta_para.add_run(f"项目：{project_name}    版本：v{version_no}")
        meta_run.font.size = Pt(11)
        meta_run.font.color.rgb = None
        meta_run.font.name = "Microsoft YaHei"

        document.add_paragraph()  # 空行分隔

        # 章节内容
        for draft in sections:
            # 章节标题（一级）
            section_title = _SECTION_NAMES.get(draft.code, draft.code)
            heading = document.add_heading(section_title, level=1)
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].bold = True

            # 章节内容行，带样式分组
            content_lines = draft.content.splitlines()
            for line in content_lines:
                if not line.strip():
                    document.add_paragraph()  # 空行保留分段
                    continue
                # 列表项（以 "- " 或 "• " 开头）
                if line.lstrip().startswith(("- ", "• ")):
                    para = document.add_paragraph(style="List Bullet")
                    para.add_run(line.lstrip()[2:])
                # 带标签行（如 "[CRITICAL] " 或 "[建议] "）
                elif line and line[0] == "[":
                    para = document.add_paragraph()
                    tag_end = line.index("]") + 1
                    tag_text = line[:tag_end]
                    body_text = line[tag_end:]
                    tag_run = para.add_run(tag_text)
                    # 根据标签着色
                    if "[CRITICAL]" in tag_text or "[严重]" in tag_text:
                        tag_run.font.color.rgb = None
                        tag_run.bold = True
                    elif "[HIGH]" in tag_text or "[高]" in tag_text:
                        tag_run.font.color.rgb = None
                        tag_run.bold = True
                    elif "[建议]" in tag_text or "[RECOMMEND]" in tag_text:
                        tag_run.font.color.rgb = None
                    para.add_run(body_text)
                else:
                    document.add_paragraph(line)

            document.add_paragraph()  # 章节后空行

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _render_pdf(project_name: str, version_no: int, sections: list[_DraftSection]) -> bytes:
        output = BytesIO()
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        document = canvas.Canvas(output, pagesize=A4)
        document.setTitle(f"投标分析报告 v{version_no} - {project_name}")
        width, height = A4
        margin = 48
        x = margin

        def new_page() -> None:
            nonlocal y
            document.showPage()
            # ``showPage`` starts a new physical page but does not restore our
            # cursor. Without this reset every subsequent line is written below
            # the page and triggers another empty page.
            y = height - margin

        def check_page(line_height: float = 16) -> None:
            nonlocal y
            if y - line_height < margin:
                new_page()

        def write_line(text: str, size: int = 10, bold: bool = False) -> None:
            nonlocal y
            check_page(size + 8)
            document.setFont("STSong-Light", size)
            document.drawString(x, y, text[:80])
            y -= size + 6

        def write_separator() -> None:
            nonlocal y
            check_page(4)
            document.setStrokeColorRGB(0.8, 0.8, 0.8)
            document.setLineWidth(0.5)
            document.line(x, y, width - margin, y)
            y -= 8

        y = height - margin

        # 封面标题
        write_line("投标综合分析报告", 20)
        y -= 4
        write_line(f"项目：{project_name}", 11)
        ts = datetime.now(UTC).strftime("%Y-%m-%d %H:%M")
        write_line(f"版本：v{version_no}    生成时间：{ts}", 10)
        write_separator()

        # 章节内容
        for draft in sections:
            check_page(24)
            section_title = _SECTION_NAMES.get(draft.code, draft.code)
            write_line(section_title, 14)
            y -= 2
            write_separator()

            for raw_line in draft.content.splitlines() or []:
                if not raw_line.strip():
                    y -= 6
                    continue
                raw_line = raw_line.lstrip("#").lstrip().replace("**", "")
                # 智能换行
                for start in range(0, len(raw_line), 55):
                    write_line(raw_line[start : start + 55])
            y -= 8  # 章节间距

        document.save()
        return output.getvalue()

    @staticmethod
    def _render_md(project_name: str, version_no: int, sections: list[_DraftSection]) -> bytes:
        lines = [
            "# 投标综合分析报告",
            "",
            f"**项目**：{project_name}",
            f"**版本**：v{version_no}",
            "",
        ]
        for draft in sections:
            lines.append(f"## {_SECTION_NAMES.get(draft.code, draft.code)}")
            lines.append("")
            for line in draft.content.splitlines():
                lines.append(line)
            lines.append("")
        return "\n".join(lines).encode("utf-8")

    def _compensate_output(self, object_key: str) -> None:
        try:
            self._storage.delete_object(object_key)
        except ObjectStorageUnavailable:
            pass

    def _verify_object_hash(self, object_key: str, expected_content: bytes) -> None:
        actual_hash = sha256()
        try:
            with self._storage.open_object(object_key) as source:
                for chunk in source.stream(amt=1024 * 1024):
                    actual_hash.update(chunk)
        except ObjectStorageUnavailable:
            raise
        if actual_hash.digest() != sha256(expected_content).digest():
            raise ObjectStorageUnavailable("Report object hash verification failed")

    @staticmethod
    def _require_writer(role_codes: set[str]) -> None:
        if not can_generate_reports(role_codes):
            raise DomainError("PERMISSION_DENIED", "无权生成报告。", 403)

    @staticmethod
    def _require_reader(role_codes: set[str]) -> None:
        if not can_read_reports(role_codes):
            raise DomainError("PERMISSION_DENIED", "无权访问报告。", 403)

    @staticmethod
    def _is_read_only(role_codes: set[str]) -> bool:
        return READ_ONLY in role_codes and not can_generate_reports(role_codes)

    def _response(self, report: Report) -> ReportResponse:
        return ReportResponse(
            id=report.id,
            project_id=report.project_id,
            version_no=report.version_no,
            status=report.status,
            error_code=report.error_code,
            error_message=report.error_message,
            generated_by=report.generated_by,
            generated_at=report.generated_at,
            created_at=report.created_at,
            sections=[
                ReportSectionResponse(
                    section_code=section.section_code,
                    order_no=section.order_no,
                    content_markdown=section.content_markdown,
                    evidence_ids=self._reports.list_evidence_ids(section.id),
                )
                for section in self._reports.list_sections(report.id)
            ],
        )

    @staticmethod
    def _task_response(task) -> TaskResponse:
        return TaskResponse(
            id=task.id,
            task_type=task.task_type,
            target_type=task.target_type,
            target_id=task.target_id,
            status=task.status,
            attempt=task.attempt,
            error_code=task.error_code,
            error_message=task.error_message,
            created_at=task.created_at,
            started_at=task.started_at,
            completed_at=task.completed_at,
        )
