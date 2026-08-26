from __future__ import annotations

import hashlib
import json
from collections.abc import Iterator
from dataclasses import dataclass
from datetime import UTC, datetime
from decimal import ROUND_HALF_UP, Decimal
from io import BytesIO
from uuid import UUID, uuid4

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import Settings
from app.core.constants import BID_SPECIALIST, LEGAL_COMPLIANCE, PROJECT_OWNER, SYSTEM_ADMIN
from app.core.errors import DomainError
from app.db.models import (
    AgentRun,
    AgentRunEvidence,
    ChallengeDraft,
    ChallengeDraftEvidence,
    CompetitiveAnalysis,
    CompetitiveAnalysisEvidence,
    CompetitiveFinding,
    CompetitiveFindingEvidence,
    CompetitiveFindingKnowledge,
    DocumentVersion,
    Evidence,
    GraphEdge,
    GraphNode,
    IntegrationRun,
    KnowledgeEntry,
    KnowledgeVersion,
    MarketCheck,
    Notification,
    ProjectComment,
    ProjectMember,
    QuoteScenario,
    SearchChunk,
    Task,
    User,
    WorkItem,
)
from app.db.repositories.advanced_repository import AdvancedRepository
from app.db.repositories.document_repository import DocumentRepository
from app.db.repositories.evidence_repository import EvidenceRepository
from app.db.repositories.requirement_repository import RequirementRepository
from app.db.repositories.risk_repository import RiskRepository
from app.integrations.ai.llm import DeepSeekV4FlashClient
from app.integrations.external_connectors import ConnectorExecutionError, ConnectorExecutor
from app.integrations.object_storage import MinioObjectStorage, ObjectStorageUnavailable
from app.integrations.task_publisher import TaskPublisher
from app.integrations.vector_store import PgVectorStore
from app.schemas.advanced import (
    AgentRunCreate,
    AgentRunResponse,
    ChallengeDraftCreate,
    ChallengeDraftResponse,
    ChallengeDraftReview,
    CompetitiveAnalysisCreate,
    CompetitiveAnalysisResponse,
    CompetitiveFindingDraft,
    CompetitiveFindingResponse,
    CompetitiveFindingReviewRequest,
    ConnectorResponse,
    ConnectorUpdate,
    GraphEdgeResponse,
    GraphNodeResponse,
    IntegrationRunCreate,
    IntegrationRunResponse,
    KnowledgeCreateRequest,
    KnowledgeResponse,
    KnowledgeRevisionRequest,
    MarketCheckCreate,
    MarketCheckResponse,
    NotificationResponse,
    ProjectCommentCreate,
    ProjectCommentResponse,
    ProjectGraphResponse,
    QuoteScenarioCreate,
    QuoteScenarioResponse,
    WorkItemCreate,
    WorkItemResponse,
    WorkItemUpdate,
)
from app.services.audit_service import AuditService
from app.services.project_service import ProjectService

_WRITER_ROLES = {SYSTEM_ADMIN, PROJECT_OWNER, BID_SPECIALIST, LEGAL_COMPLIANCE}
_FINDING_CATEGORIES = {
    "BRAND_OR_PARAMETER": ("品牌", "型号", "参数", "专利", "指定"),
    "EXCESSIVE_QUALIFICATION": ("注册资本", "资质等级", "业绩", "年限"),
    "GEOGRAPHIC_RESTRICTION": ("本地", "当地", "注册地", "属地"),
    "UNIQUE_SUPPLY": ("唯一", "独家", "指定供应", "特定供应商"),
    "INCONSISTENT_REQUIREMENT": ("不一致", "矛盾", "同时满足"),
}
_FINDING_TITLES = {
    "BRAND_OR_PARAMETER": "可能存在品牌或特定参数限制，建议法务核查",
    "EXCESSIVE_QUALIFICATION": "可能存在较高资格门槛，建议法务核查",
    "GEOGRAPHIC_RESTRICTION": "可能存在地域性限制，建议法务核查",
    "UNIQUE_SUPPLY": "可能存在唯一来源或指定供应限制，建议法务核查",
    "INCONSISTENT_REQUIREMENT": "条款可能存在不一致表述，建议法务核查",
}
_TARGET_MODELS = {
    "COMPETITIVE_FINDING": CompetitiveFinding,
    "CHALLENGE_DRAFT": ChallengeDraft,
    "QUOTE_SCENARIO": QuoteScenario,
    "WORK_ITEM": WorkItem,
    "MARKET_CHECK": MarketCheck,
    "AGENT_RUN": AgentRun,
}


@dataclass(frozen=True, slots=True)
class AuthorizedAdvancedDownload:
    file_name: str
    mime_type: str
    stream: Iterator[bytes]


class AdvancedService:
    def __init__(
        self,
        session: Session,
        storage: MinioObjectStorage,
        vector_store: PgVectorStore,
        settings: Settings,
        publisher: TaskPublisher | None = None,
    ) -> None:
        self._session = session
        self._storage = storage
        self._vector_store = vector_store
        self._settings = settings
        self._publisher = publisher
        self._repo = AdvancedRepository(session)
        self._projects = ProjectService(session)
        self._documents = DocumentRepository(session)
        self._evidences = EvidenceRepository(session)
        self._requirements = RequirementRepository(session)
        self._risks = RiskRepository(session)
        self._audit = AuditService(session)
        self._llm = DeepSeekV4FlashClient(settings) if settings.ai_is_configured else None

    # Knowledge base
    def create_knowledge(
        self, actor_id: UUID, role_codes: set[str], payload: KnowledgeCreateRequest
    ) -> KnowledgeResponse:
        self._require_knowledge_manager(role_codes)
        now = datetime.now(UTC)
        entry = KnowledgeEntry(
            id=uuid4(),
            knowledge_type=payload.knowledge_type,
            title=payload.title,
            authority=payload.authority,
            source_reference=payload.source_reference,
            created_at=now,
            created_by=actor_id,
            updated_at=now,
            deleted_at=None,
        )
        version = KnowledgeVersion(
            id=uuid4(),
            knowledge_entry_id=entry.id,
            version_no=1,
            status="DRAFT",
            content=payload.content,
            issued_on=payload.issued_on,
            effective_on=payload.effective_on,
            citation_note=payload.citation_note,
            source_document_version_id=None,
            source_evidence_id=None,
            published_at=None,
            published_by=None,
            created_at=now,
            created_by=actor_id,
        )
        evidence = self._manual_knowledge_evidence(entry, version, actor_id)
        # `knowledge_versions.source_evidence_id` is a database FK.  Flush the
        # entry and its synthetic evidence first, otherwise SQLAlchemy may
        # insert the version before the evidence because there is no ORM
        # relationship between them.
        self._repo.add_all([entry, evidence])
        self._session.flush()
        version.source_evidence_id = evidence.id
        self._repo.add(version)
        self._audit.record(
            actor_id=actor_id,
            action="CREATE_KNOWLEDGE_ENTRY",
            target_type="KNOWLEDGE_ENTRY",
            target_id=entry.id,
            after={"knowledge_type": entry.knowledge_type, "version_no": 1},
        )
        self._session.commit()
        return self._knowledge_response(entry, version)

    def list_knowledge(
        self, actor_id: UUID, role_codes: set[str], query: str | None
    ) -> list[KnowledgeResponse]:
        del actor_id
        can_manage = self._can_manage_knowledge(role_codes)
        return [
            self._knowledge_response(entry, version)
            for entry, version in self._repo.list_knowledge(
                published_only=not can_manage, query=query
            )
        ]

    def delete_knowledge(
        self, entry_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> None:
        self._require_knowledge_manager(role_codes)
        entry = self._repo.get_knowledge_entry(entry_id)
        if not entry or entry.deleted_at is not None:
            raise DomainError("RESOURCE_NOT_FOUND", "知识条目不存在", 404)

        # 获取所有版本信息（用于收集外部存储的 PK）
        versions = self._session.scalars(
            select(KnowledgeVersion).where(KnowledgeVersion.knowledge_entry_id == entry_id)
        ).all()

        evidence_ids: list[UUID] = []
        source_document_version_ids: list[UUID] = []
        chunk_pks: list[str] = []
        object_keys_to_delete: list[str] = []

        for version in versions:
            if version.source_evidence_id:
                evidence_ids.append(version.source_evidence_id)
            if version.source_document_version_id:
                source_document_version_ids.append(version.source_document_version_id)
            object_keys_to_delete.append(f"knowledge-source/{entry.id}/{version.id}/source")

        # 收集需要清除向量的 chunk PK（通过 evidence 关联）
        all_evidence_ids: set[UUID] = set(evidence_ids)
        if source_document_version_ids:
            doc_version_evids = list(self._session.scalars(
                select(Evidence.id).where(
                    Evidence.document_version_id.in_(source_document_version_ids)
                )
            ).all())
            all_evidence_ids.update(doc_version_evids)
        if all_evidence_ids:
            chunks = self._session.scalars(
                select(SearchChunk).where(SearchChunk.evidence_id.in_(all_evidence_ids))
            ).all()
            for chunk in chunks:
                chunk_pks.append(str(chunk.id))

        # 先删外部存储，失败则整个事务回滚
        if chunk_pks:
            self._vector_store.delete(chunk_pks)
        for object_key in object_keys_to_delete:
            self._storage.delete_object(object_key)

        # 先删 DocumentVersion（通过 CASCADE 链删除 DocumentNode、Evidence 等junction表），
        # 再删 KnowledgeEntry（通过 CASCADE 链删除 KnowledgeVersion）。
        # 两步之间无 FK 依赖，因为删 DocumentVersion 不影响 KnowledgeVersion。
        if source_document_version_ids:
            # 显式删除 Task（多态关联，无 FK 约束），CASCADE 到 AiRun
            self._session.execute(
                Task.__table__.delete().where(
                    Task.target_type == "DOCUMENT_VERSION",
                    Task.target_id.in_(source_document_version_ids),
                )
            )
            self._session.execute(
                DocumentVersion.__table__.delete().where(
                    DocumentVersion.id.in_(source_document_version_ids)
                )
            )
        self._repo.delete_knowledge_entry(entry_id)

        self._audit.record(
            actor_id=actor_id,
            action="DELETE_KNOWLEDGE_ENTRY",
            target_type="KNOWLEDGE_ENTRY",
            target_id=entry_id,
        )

    def revise_knowledge(
        self,
        entry_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: KnowledgeRevisionRequest,
    ) -> KnowledgeResponse:
        self._require_knowledge_manager(role_codes)
        entry = self._repo.get_knowledge_entry(entry_id, for_update=True)
        if entry is None or entry.deleted_at is not None:
            raise DomainError("RESOURCE_NOT_FOUND", "知识条目不存在", 404)
        now = datetime.now(UTC)
        version = KnowledgeVersion(
            id=uuid4(),
            knowledge_entry_id=entry.id,
            version_no=self._repo.next_knowledge_version(entry.id),
            status="DRAFT",
            content=payload.content,
            issued_on=payload.issued_on,
            effective_on=payload.effective_on,
            citation_note=payload.citation_note,
            source_document_version_id=None,
            source_evidence_id=None,
            published_at=None,
            published_by=None,
            created_at=now,
            created_by=actor_id,
        )
        evidence = self._manual_knowledge_evidence(entry, version, actor_id)
        entry.updated_at = now
        self._repo.add(evidence)
        self._session.flush()
        version.source_evidence_id = evidence.id
        self._repo.add(version)
        self._audit.record(
            actor_id=actor_id,
            action="REVISE_KNOWLEDGE_ENTRY",
            target_type="KNOWLEDGE_ENTRY",
            target_id=entry.id,
            after={"version_no": version.version_no},
        )
        self._session.commit()
        return self._knowledge_response(entry, version)

    def publish_knowledge(
        self, version_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> KnowledgeResponse:
        self._require_knowledge_manager(role_codes)
        version = self._repo.get_knowledge_version(version_id)
        if version is None:
            raise DomainError("RESOURCE_NOT_FOUND", "知识版本不存在", 404)
        entry = self._repo.get_knowledge_entry(version.knowledge_entry_id, for_update=True)
        if entry is None or entry.deleted_at is not None:
            raise DomainError("RESOURCE_NOT_FOUND", "知识条目不存在", 404)
        if version.status != "DRAFT":
            raise DomainError("INVALID_STATE_TRANSITION", "仅草稿版本可以发布", 409)
        if not version.content.strip():
            raise DomainError("KNOWLEDGE_VERSION_NOT_READY", "知识版本尚未生成可发布正文", 409)
        if version.source_document_version_id is not None:
            source_version = self._documents.get_version(version.source_document_version_id)
            if source_version is None or source_version.parse_status != "READY":
                raise DomainError(
                    "KNOWLEDGE_VERSION_NOT_READY",
                    "源文件尚未完成解析和向量索引，不能发布",
                    409,
                )
        for prior in self._session.scalars(
            select(KnowledgeVersion).where(
                KnowledgeVersion.knowledge_entry_id == entry.id,
                KnowledgeVersion.status == "PUBLISHED",
            )
        ):
            prior.status = "DRAFT"
        now = datetime.now(UTC)
        version.status = "PUBLISHED"
        version.published_at = now
        version.published_by = actor_id
        entry.updated_at = now
        self._audit.record(
            actor_id=actor_id,
            action="PUBLISH_KNOWLEDGE_VERSION",
            target_type="KNOWLEDGE_VERSION",
            target_id=version.id,
            after={"entry_id": str(entry.id), "version_no": version.version_no},
        )
        self._session.commit()
        return self._knowledge_response(entry, version)

    def unpublish_knowledge(
        self, version_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> KnowledgeResponse:
        self._require_knowledge_manager(role_codes)
        version = self._repo.get_knowledge_version(version_id)
        if version is None:
            raise DomainError("RESOURCE_NOT_FOUND", "知识版本不存在", 404)
        entry = self._repo.get_knowledge_entry(version.knowledge_entry_id, for_update=True)
        if entry is None or entry.deleted_at is not None:
            raise DomainError("RESOURCE_NOT_FOUND", "知识条目不存在", 404)
        if version.status != "PUBLISHED":
            raise DomainError("INVALID_STATE_TRANSITION", "仅已发布版本可以停用", 409)
        version.status = "DRAFT"
        version.published_at = None
        version.published_by = None
        entry.updated_at = datetime.now(UTC)
        self._audit.record(
            actor_id=actor_id,
            action="UNPUBLISH_KNOWLEDGE_VERSION",
            target_type="KNOWLEDGE_VERSION",
            target_id=version.id,
            after={"entry_id": str(entry.id), "version_no": version.version_no},
        )
        self._session.commit()
        return self._knowledge_response(entry, version)

    # Competitive terms
    def create_competitive_analysis(
        self,
        project_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: CompetitiveAnalysisCreate,
    ) -> CompetitiveAnalysisResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        evidence_ids = self._validated_project_evidence_ids(project_id, payload.evidence_ids)
        if payload.requirement_id is not None:
            requirement = self._requirements.get(payload.requirement_id)
            if requirement is None or requirement.project_id != project_id:
                raise DomainError("RESOURCE_NOT_FOUND", "Requirement 不存在", 404)
        knowledge_ids = self._published_knowledge_ids(payload.knowledge_version_ids)
        now = datetime.now(UTC)
        analysis = CompetitiveAnalysis(
            id=uuid4(),
            project_id=project_id,
            requirement_id=payload.requirement_id,
            status="RUNNING",
            method=payload.method,
            summary=None,
            created_at=now,
            created_by=actor_id,
            completed_at=None,
        )
        self._repo.add(analysis)
        self._repo.add_all(
            [
                CompetitiveAnalysisEvidence(analysis_id=analysis.id, evidence_id=evidence_id)
                for evidence_id in evidence_ids
            ]
        )
        texts = [self._evidences.get(evidence_id).quoted_text or "" for evidence_id in evidence_ids]

        if payload.method == "LLM_ANALYSIS":
            findings = self._llm_findings(analysis.id, texts, evidence_ids, knowledge_ids, now)
        elif payload.method == "HYBRID":
            keyword = self._keyword_findings(analysis.id, texts, evidence_ids, knowledge_ids, now)
            llm = self._llm_findings(analysis.id, texts, evidence_ids, knowledge_ids, now)
            findings = self._deduplicate_findings(keyword, llm)
        else:
            findings = self._keyword_findings(analysis.id, texts, evidence_ids, knowledge_ids, now)

        self._repo.add_all(findings)
        analysis.status = "READY"
        analysis.completed_at = now
        method_label = {"DETERMINISTIC_RULES": "规则化", "LLM_ANALYSIS": "LLM", "HYBRID": "混合"}
        analysis.summary = (
            f"已完成{method_label.get(payload.method, payload.method)}竞争性条款核查，"
            f"生成 {len(findings)} 条风险提示；所有结果均需人工法务复核。"
        )
        self._audit.record(
            actor_id=actor_id,
            action="RUN_COMPETITIVE_ANALYSIS",
            target_type="COMPETITIVE_ANALYSIS",
            target_id=analysis.id,
            project_id=project_id,
            after={"finding_count": len(findings), "method": analysis.method},
        )
        self._session.commit()
        return self._analysis_response(analysis)

    def list_competitive_analyses(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[CompetitiveAnalysisResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [
            self._analysis_response(analysis) for analysis in self._repo.list_analyses(project_id)
        ]

    def review_competitive_finding(
        self,
        project_id: UUID,
        finding_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: CompetitiveFindingReviewRequest,
    ) -> CompetitiveFindingResponse:
        self._require_legal_reviewer(role_codes)
        finding = self._repo.get_finding(finding_id, for_update=True)
        if finding is None:
            raise DomainError("RESOURCE_NOT_FOUND", "竞争性条款发现不存在", 404)
        analysis = self._repo.get_analysis(finding.analysis_id)
        if analysis is None or analysis.project_id != project_id:
            raise DomainError("RESOURCE_NOT_FOUND", "竞争性条款发现不存在", 404)
        self._projects.get_visible(project_id, actor_id, role_codes)
        if payload.status != "PENDING" and not payload.resolution:
            raise DomainError("VALIDATION_ERROR", "复核状态变更必须填写说明", 422)
        knowledge_ids = self._published_knowledge_ids(payload.knowledge_version_ids)
        now = datetime.now(UTC)
        finding.status = payload.status
        finding.resolution = payload.resolution
        finding.reviewed_at = now
        finding.reviewed_by = actor_id
        existing = self._finding_knowledge_ids(finding.id)
        self._repo.add_all(
            [
                CompetitiveFindingKnowledge(finding_id=finding.id, knowledge_version_id=version_id)
                for version_id in knowledge_ids
                if version_id not in existing
            ]
        )
        self._audit.record(
            actor_id=actor_id,
            action="REVIEW_COMPETITIVE_FINDING",
            target_type="COMPETITIVE_FINDING",
            target_id=finding.id,
            project_id=project_id,
            after={"status": finding.status},
        )
        self._session.commit()
        return self._finding_response(finding)

    # Challenge drafts
    def create_challenge_draft(
        self,
        project_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: ChallengeDraftCreate,
    ) -> ChallengeDraftResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        evidence_ids = self._validated_project_evidence_ids(project_id, payload.evidence_ids)
        now = datetime.now(UTC)
        draft = ChallengeDraft(
            id=uuid4(),
            project_id=project_id,
            title=payload.title,
            subject=payload.subject,
            fact_statement=payload.fact_statement,
            requested_action=payload.requested_action,
            status="DRAFT",
            review_note=None,
            reviewed_by=None,
            reviewed_at=None,
            docx_object_key=None,
            pdf_object_key=None,
            created_at=now,
            created_by=actor_id,
            updated_at=now,
        )
        self._repo.add(draft)
        self._repo.add_all(
            [
                ChallengeDraftEvidence(challenge_draft_id=draft.id, evidence_id=evidence_id)
                for evidence_id in evidence_ids
            ]
        )
        self._session.flush()
        docx_key = f"challenge-drafts/{project_id}/{draft.id}/draft.docx"
        pdf_key = f"challenge-drafts/{project_id}/{draft.id}/draft.pdf"
        try:
            self._storage.put_bytes(
                docx_key,
                self._render_challenge_docx(draft, evidence_ids),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self._storage.put_bytes(
                pdf_key, self._render_challenge_pdf(draft, evidence_ids), "application/pdf"
            )
        except ObjectStorageUnavailable as exc:
            self._session.rollback()
            self._compensate_object(docx_key)
            self._compensate_object(pdf_key)
            raise DomainError(
                "OBJECT_STORAGE_UNAVAILABLE", "质疑函草稿对象存储不可用", 503
            ) from exc
        draft.docx_object_key = docx_key
        draft.pdf_object_key = pdf_key
        self._audit.record(
            actor_id=actor_id,
            action="CREATE_CHALLENGE_DRAFT",
            target_type="CHALLENGE_DRAFT",
            target_id=draft.id,
            project_id=project_id,
            after={"evidence_count": len(evidence_ids)},
        )
        self._session.commit()
        return self._challenge_response(draft)

    def list_challenge_drafts(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[ChallengeDraftResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [self._challenge_response(item) for item in self._repo.list_challenges(project_id)]

    def review_challenge_draft(
        self,
        project_id: UUID,
        draft_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: ChallengeDraftReview,
    ) -> ChallengeDraftResponse:
        self._require_legal_reviewer(role_codes)
        draft = self._repo.get_challenge(draft_id, for_update=True)
        if draft is None or draft.project_id != project_id:
            raise DomainError("RESOURCE_NOT_FOUND", "质疑函草稿不存在", 404)
        self._projects.get_visible(project_id, actor_id, role_codes)
        if draft.status in {"APPROVED", "REJECTED"}:
            raise DomainError("INVALID_STATE_TRANSITION", "已终审草稿不能再次审核", 409)
        now = datetime.now(UTC)
        draft.status = payload.status
        draft.review_note = payload.review_note
        draft.reviewed_by = actor_id
        draft.reviewed_at = now
        draft.updated_at = now
        self._audit.record(
            actor_id=actor_id,
            action="REVIEW_CHALLENGE_DRAFT",
            target_type="CHALLENGE_DRAFT",
            target_id=draft.id,
            project_id=project_id,
            after={"status": draft.status},
        )
        self._session.commit()
        return self._challenge_response(draft)

    def create_challenge_download(
        self,
        draft_id: UUID,
        draft_format: str,
        actor_id: UUID,
        role_codes: set[str],
    ) -> AuthorizedAdvancedDownload:
        draft = self._repo.get_challenge(draft_id)
        if draft is None:
            raise DomainError("RESOURCE_NOT_FOUND", "质疑函草稿不存在", 404)
        self._projects.get_visible(draft.project_id, actor_id, role_codes)
        object_key = draft.docx_object_key if draft_format == "docx" else draft.pdf_object_key
        if object_key is None:
            raise DomainError("RESOURCE_NOT_FOUND", "草稿产物不存在", 404)
        self._audit.record(
            actor_id=actor_id,
            action="DOWNLOAD_CHALLENGE_DRAFT",
            target_type="CHALLENGE_DRAFT",
            target_id=draft.id,
            project_id=draft.project_id,
            after={"format": draft_format},
        )
        self._session.commit()

        def iterator() -> Iterator[bytes]:
            with self._storage.open_object(object_key) as source:
                yield from source.stream(amt=1024 * 1024)

        return AuthorizedAdvancedDownload(
            file_name=f"challenge-draft-{draft.id}.{draft_format}",
            mime_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if draft_format == "docx"
                else "application/pdf"
            ),
            stream=iterator(),
        )

    # Quote sandbox
    def create_quote_scenario(
        self,
        project_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: QuoteScenarioCreate,
    ) -> QuoteScenarioResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        now = datetime.now(UTC)
        scenario = QuoteScenario(
            id=uuid4(),
            project_id=project_id,
            parent_scenario_id=None,
            name=payload.name,
            version_no=self._repo.next_quote_version(project_id, payload.name),
            status="DRAFT",
            cost_excluding_tax=payload.cost_excluding_tax,
            tax_rate=payload.tax_rate,
            target_margin_rate=payload.target_margin_rate,
            risk_adjustment=payload.risk_adjustment,
            expected_score=payload.expected_score,
            assumptions=payload.assumptions,
            calculations=self._quote_calculations(payload),
            locked_at=None,
            locked_by=None,
            created_at=now,
            created_by=actor_id,
            updated_at=now,
        )
        self._repo.add(scenario)
        self._audit.record(
            actor_id=actor_id,
            action="CREATE_QUOTE_SCENARIO",
            target_type="QUOTE_SCENARIO",
            target_id=scenario.id,
            project_id=project_id,
            after={"version_no": scenario.version_no},
        )
        self._session.commit()
        return self._quote_response(scenario)

    def list_quote_scenarios(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[QuoteScenarioResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [self._quote_response(item) for item in self._repo.list_quotes(project_id)]

    def lock_quote_scenario(
        self, project_id: UUID, scenario_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> QuoteScenarioResponse:
        project = self._projects.get_visible(project_id, actor_id, role_codes)
        self._projects.require_writable(project)
        if SYSTEM_ADMIN not in role_codes and project.owner_id != actor_id:
            raise DomainError("PERMISSION_DENIED", "仅项目负责人可锁定报价情景", 403)
        scenario = self._repo.get_quote(scenario_id, for_update=True)
        if scenario is None or scenario.project_id != project_id:
            raise DomainError("RESOURCE_NOT_FOUND", "报价情景不存在", 404)
        if scenario.status != "DRAFT":
            raise DomainError("INVALID_STATE_TRANSITION", "仅草稿情景可以锁定", 409)
        now = datetime.now(UTC)
        scenario.status = "LOCKED"
        scenario.locked_at = now
        scenario.locked_by = actor_id
        scenario.updated_at = now
        self._audit.record(
            actor_id=actor_id,
            action="LOCK_QUOTE_SCENARIO",
            target_type="QUOTE_SCENARIO",
            target_id=scenario.id,
            project_id=project_id,
        )
        self._session.commit()
        return self._quote_response(scenario)

    def copy_quote_scenario(
        self, project_id: UUID, scenario_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> QuoteScenarioResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        source = self._repo.get_quote(scenario_id)
        if source is None or source.project_id != project_id:
            raise DomainError("RESOURCE_NOT_FOUND", "报价情景不存在", 404)
        now = datetime.now(UTC)
        copied = QuoteScenario(
            id=uuid4(),
            project_id=project_id,
            parent_scenario_id=source.id,
            name=source.name,
            version_no=self._repo.next_quote_version(project_id, source.name),
            status="DRAFT",
            cost_excluding_tax=source.cost_excluding_tax,
            tax_rate=source.tax_rate,
            target_margin_rate=source.target_margin_rate,
            risk_adjustment=source.risk_adjustment,
            expected_score=source.expected_score,
            assumptions=dict(source.assumptions),
            calculations=dict(source.calculations),
            locked_at=None,
            locked_by=None,
            created_at=now,
            created_by=actor_id,
            updated_at=now,
        )
        self._repo.add(copied)
        self._audit.record(
            actor_id=actor_id,
            action="COPY_QUOTE_SCENARIO",
            target_type="QUOTE_SCENARIO",
            target_id=copied.id,
            project_id=project_id,
            after={"parent_scenario_id": str(source.id)},
        )
        self._session.commit()
        return self._quote_response(copied)

    # Collaboration
    def create_comment(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str], payload: ProjectCommentCreate
    ) -> ProjectCommentResponse:
        self._projects.get_visible(project_id, actor_id, role_codes)
        self._validate_target(project_id, payload.target_type, payload.target_id)
        comment = ProjectComment(
            id=uuid4(),
            project_id=project_id,
            target_type=payload.target_type,
            target_id=payload.target_id,
            content=payload.content,
            created_at=datetime.now(UTC),
            created_by=actor_id,
        )
        self._repo.add(comment)
        self._audit.record(
            actor_id=actor_id,
            action="CREATE_PROJECT_COMMENT",
            target_type="PROJECT_COMMENT",
            target_id=comment.id,
            project_id=project_id,
        )
        self._session.commit()
        return self._comment_response(comment)

    def list_comments(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[ProjectCommentResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [self._comment_response(item) for item in self._repo.list_comments(project_id)]

    def create_work_item(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str], payload: WorkItemCreate
    ) -> WorkItemResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        self._validate_target(project_id, payload.target_type, payload.target_id)
        if payload.assignee_id is not None and not self._is_project_assignee(
            project_id, payload.assignee_id
        ):
            raise DomainError("VALIDATION_ERROR", "待办只能指派给项目成员", 422)
        now = datetime.now(UTC)
        item = WorkItem(
            id=uuid4(),
            project_id=project_id,
            title=payload.title,
            description=payload.description,
            status="OPEN",
            assignee_id=payload.assignee_id,
            due_at=payload.due_at,
            target_type=payload.target_type,
            target_id=payload.target_id,
            closing_note=None,
            created_at=now,
            created_by=actor_id,
            updated_at=now,
        )
        self._repo.add(item)
        if item.assignee_id is not None and item.assignee_id != actor_id:
            self._repo.add(
                Notification(
                    id=uuid4(),
                    user_id=item.assignee_id,
                    project_id=project_id,
                    notification_type="WORK_ITEM_ASSIGNED",
                    payload={"work_item_id": str(item.id), "title": item.title},
                    read_at=None,
                    created_at=now,
                )
            )
        self._audit.record(
            actor_id=actor_id,
            action="CREATE_WORK_ITEM",
            target_type="WORK_ITEM",
            target_id=item.id,
            project_id=project_id,
            after={"assignee_id": str(item.assignee_id) if item.assignee_id else None},
        )
        self._session.commit()
        return self._work_item_response(item)

    def list_work_items(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[WorkItemResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [self._work_item_response(item) for item in self._repo.list_work_items(project_id)]

    def update_work_item(
        self,
        project_id: UUID,
        work_item_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: WorkItemUpdate,
    ) -> WorkItemResponse:
        project = self._projects.get_visible(project_id, actor_id, role_codes)
        self._projects.require_writable(project)
        item = self._repo.get_work_item(work_item_id, for_update=True)
        if item is None or item.project_id != project_id:
            raise DomainError("RESOURCE_NOT_FOUND", "待办不存在", 404)
        permitted = (
            SYSTEM_ADMIN in role_codes
            or project.owner_id == actor_id
            or item.assignee_id == actor_id
        )
        if not permitted:
            raise DomainError("PERMISSION_DENIED", "无权更新该待办", 403)
        if payload.status in {"DONE", "CANCELLED"} and not payload.closing_note:
            raise DomainError("VALIDATION_ERROR", "关闭或取消待办必须填写说明", 422)
        item.status = payload.status
        item.closing_note = payload.closing_note
        item.updated_at = datetime.now(UTC)
        self._audit.record(
            actor_id=actor_id,
            action="UPDATE_WORK_ITEM",
            target_type="WORK_ITEM",
            target_id=item.id,
            project_id=project_id,
            after={"status": item.status},
        )
        self._session.commit()
        return self._work_item_response(item)

    def list_notifications(self, actor_id: UUID) -> list[NotificationResponse]:
        return [
            self._notification_response(item) for item in self._repo.list_notifications(actor_id)
        ]

    def mark_notification_read(self, notification_id: UUID, actor_id: UUID) -> NotificationResponse:
        notification = self._repo.get_notification(notification_id, actor_id)
        if notification is None:
            raise DomainError("RESOURCE_NOT_FOUND", "通知不存在", 404)
        if notification.read_at is None:
            notification.read_at = datetime.now(UTC)
            self._session.commit()
        return self._notification_response(notification)

    # P2 market check, graph, agent and connector control
    def create_market_check(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str], payload: MarketCheckCreate
    ) -> MarketCheckResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        if (payload.requirement_id is None) == (payload.evidence_id is None):
            raise DomainError(
                "VALIDATION_ERROR", "市场核查必须关联一个 Requirement 或 Evidence", 422
            )
        if payload.requirement_id is not None:
            requirement = self._requirements.get(payload.requirement_id)
            if requirement is None or requirement.project_id != project_id:
                raise DomainError("RESOURCE_NOT_FOUND", "Requirement 不存在", 404)
        if payload.evidence_id is not None:
            self._validated_project_evidence_ids(project_id, [payload.evidence_id])
        check = MarketCheck(
            id=uuid4(),
            project_id=project_id,
            requirement_id=payload.requirement_id,
            evidence_id=payload.evidence_id,
            parameter=payload.parameter,
            source_name=payload.source_name,
            source_reference=payload.source_reference,
            excerpt=payload.excerpt,
            conclusion=payload.conclusion,
            note=payload.note,
            created_at=datetime.now(UTC),
            created_by=actor_id,
        )
        self._repo.add(check)
        self._audit.record(
            actor_id=actor_id,
            action="CREATE_MARKET_CHECK",
            target_type="MARKET_CHECK",
            target_id=check.id,
            project_id=project_id,
            after={"conclusion": check.conclusion},
        )
        self._session.commit()
        return self._market_check_response(check)

    def list_market_checks(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[MarketCheckResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [
            self._market_check_response(item) for item in self._repo.list_market_checks(project_id)
        ]

    def rebuild_graph(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> ProjectGraphResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        now = datetime.now(UTC)
        project_node = self._upsert_graph_node(
            project_id, "PROJECT", str(project_id), "项目", {}, None, now
        )
        for requirement in self._requirements.list_confirmed_for_project(project_id):
            evidence_ids = self._requirements.list_evidence_ids(requirement.id)
            evidence_id = evidence_ids[0] if evidence_ids else None
            node = self._upsert_graph_node(
                project_id,
                "REQUIREMENT",
                str(requirement.id),
                requirement.title,
                {"category": requirement.category, "mandatory": requirement.is_mandatory},
                evidence_id,
                now,
            )
            self._upsert_graph_edge(
                project_id, project_node.id, node.id, "HAS_REQUIREMENT", evidence_id, actor_id, now
            )
        for risk in self._risks.list_current_for_project(project_id):
            node = self._upsert_graph_node(
                project_id,
                "RISK",
                str(risk.id),
                risk.title,
                {"severity": risk.severity, "status": risk.status},
                risk.primary_evidence_id,
                now,
            )
            self._upsert_graph_edge(
                project_id,
                project_node.id,
                node.id,
                "HAS_RISK",
                risk.primary_evidence_id,
                actor_id,
                now,
            )
        self._audit.record(
            actor_id=actor_id,
            action="REBUILD_PROJECT_GRAPH",
            target_type="PROJECT",
            target_id=project_id,
            project_id=project_id,
        )
        self._session.commit()
        return self._graph_response(project_id)

    def get_graph(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> ProjectGraphResponse:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return self._graph_response(project_id)

    def create_agent_run(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str], payload: AgentRunCreate
    ) -> AgentRunResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        requirements = self._requirements.list_confirmed_for_project(project_id)
        risks = self._risks.list_current_for_project(project_id)
        evidence_ids = list(
            dict.fromkeys(
                [
                    evidence_id
                    for requirement in requirements
                    for evidence_id in self._requirements.list_evidence_ids(requirement.id)
                ]
                + [
                    risk.primary_evidence_id
                    for risk in risks
                    if risk.primary_evidence_id is not None
                ]
            )
        )
        snapshot = {
            "workflow": payload.workflow,
            "goal": payload.goal,
            "requirements": [str(item.id) for item in requirements],
            "risks": [str(item.id) for item in risks],
            "evidence_ids": [str(item) for item in evidence_ids],
        }
        now = datetime.now(UTC)
        run = AgentRun(
            id=uuid4(),
            project_id=project_id,
            source_document_version_id=None,
            workflow=payload.workflow,
            status="RUNNING",
            goal=payload.goal,
            input_hash=self._hash(snapshot),
            thread_id=None,
            checkpoint_version=0,
            requires_human_review=False,
            result={},
            error_code=None,
            error_message=None,
            created_at=now,
            started_at=now,
            completed_at=None,
            created_by=actor_id,
        )
        self._repo.add(run)
        self._repo.add_all(
            [
                AgentRunEvidence(agent_run_id=run.id, evidence_id=evidence_id)
                for evidence_id in evidence_ids
            ]
        )
        run.result = self._agent_result(payload.workflow, requirements, risks, evidence_ids)
        run.status = "SUCCEEDED"
        run.completed_at = datetime.now(UTC)
        self._audit.record(
            actor_id=actor_id,
            action="RUN_PROJECT_AGENT",
            target_type="AGENT_RUN",
            target_id=run.id,
            project_id=project_id,
            after={"workflow": run.workflow, "mode": run.result["execution_mode"]},
        )
        self._session.commit()
        return self._agent_run_response(run)

    def list_agent_runs(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[AgentRunResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [self._agent_run_response(item) for item in self._repo.list_agent_runs(project_id)]

    def list_connectors(self, actor_id: UUID, role_codes: set[str]) -> list[ConnectorResponse]:
        del actor_id
        self._require_system_admin(role_codes)
        return [self._connector_response(item) for item in self._repo.list_connectors()]

    def update_connector(
        self, code: str, actor_id: UUID, role_codes: set[str], payload: ConnectorUpdate
    ) -> ConnectorResponse:
        self._require_system_admin(role_codes)
        connector = self._repo.get_connector(code, for_update=True)
        if connector is None:
            raise DomainError("RESOURCE_NOT_FOUND", "连接器不存在", 404)
        connector.is_enabled = payload.is_enabled
        connector.updated_at = datetime.now(UTC)
        self._audit.record(
            actor_id=actor_id,
            action="UPDATE_INTEGRATION_CONNECTOR",
            target_type="INTEGRATION_CONNECTOR",
            target_id=None,
            after={"code": code, "is_enabled": connector.is_enabled},
        )
        self._session.commit()
        return self._connector_response(connector)

    def create_integration_run(
        self,
        project_id: UUID,
        actor_id: UUID,
        role_codes: set[str],
        payload: IntegrationRunCreate,
    ) -> IntegrationRunResponse:
        self._require_project_writer(project_id, actor_id, role_codes)
        connector = self._repo.get_connector(payload.connector_code)
        if connector is None:
            raise DomainError("RESOURCE_NOT_FOUND", "连接器不存在", 404)
        now = datetime.now(UTC)
        configured = self._connector_is_configured(connector.code)
        allowed = (
            connector.is_enabled and payload.operation in connector.capabilities and configured
        )
        run = IntegrationRun(
            id=uuid4(),
            project_id=project_id,
            connector_code=connector.code,
            operation=payload.operation,
            status="QUEUED" if allowed else "FAILED",
            input_hash=self._hash(payload.payload),
            result_summary={},
            external_reference=None,
            error_code=None if allowed else "INTEGRATION_UNAVAILABLE",
            error_message=None if allowed else "连接器未启用或未完成部署配置，未发出外部请求。",
            created_at=now,
            started_at=None,
            completed_at=None if allowed else now,
            created_by=actor_id,
        )
        self._repo.add(run)
        self._audit.record(
            actor_id=actor_id,
            action="REQUEST_INTEGRATION_RUN",
            target_type="INTEGRATION_RUN",
            target_id=run.id,
            project_id=project_id,
            after={
                "connector_code": run.connector_code,
                "operation": run.operation,
                "status": run.status,
            },
        )
        self._session.commit()
        if allowed:
            if self._publisher is None:
                self._fail_integration_run(
                    run,
                    "TASK_QUEUE_UNAVAILABLE",
                    "连接器任务队列不可用，未发送外部请求。",
                )
            else:
                try:
                    self._publisher.publish_integration_run(run.id, project_id, payload.payload)
                except Exception:
                    self._session.rollback()
                    self._fail_integration_run(
                        run,
                        "TASK_QUEUE_UNAVAILABLE",
                        "连接器任务队列不可用，未发送外部请求。",
                    )
        return self._integration_run_response(run)

    def list_integration_runs(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> list[IntegrationRunResponse]:
        self._projects.get_visible(project_id, actor_id, role_codes)
        return [
            self._integration_run_response(item)
            for item in self._repo.list_integration_runs(project_id)
        ]

    def execute_integration_run(
        self,
        integration_run_id: UUID,
        project_id: UUID,
        payload: dict[str, object],
        executor: ConnectorExecutor,
    ) -> None:
        """Worker-only execution with a second authorization/configuration gate."""
        run = self._repo.get_integration_run(integration_run_id, for_update=True)
        if run is None or run.status != "QUEUED":
            self._session.rollback()
            return
        if run.project_id != project_id or run.input_hash != self._hash(payload):
            self._fail_integration_run(
                run,
                "INTEGRATION_INPUT_INVALID",
                "连接器任务参数校验失败，未发送外部请求。",
            )
            return
        connector = self._repo.get_connector(run.connector_code)
        if (
            connector is None
            or not connector.is_enabled
            or run.operation not in connector.capabilities
            or not self._connector_is_configured(run.connector_code)
        ):
            self._fail_integration_run(
                run,
                "INTEGRATION_UNAVAILABLE",
                "连接器未启用或未完成部署配置，未发送外部请求。",
            )
            return
        run.status = "RUNNING"
        run.started_at = datetime.now(UTC)
        run.error_code = None
        run.error_message = None
        self._session.commit()
        try:
            result = executor.execute(
                run.connector_code, run.operation, run.id, run.project_id, payload
            )
        except ConnectorExecutionError:
            self._fail_integration_run(
                run,
                "INTEGRATION_REQUEST_FAILED",
                "外部连接器请求失败，请核查部署端点与授权后重试。",
            )
            return
        except Exception:
            self._fail_integration_run(
                run,
                "INTEGRATION_EXECUTION_FAILED",
                "外部连接器执行失败，未保存外部响应内容。",
            )
            return
        run.status = "SUCCEEDED"
        run.result_summary = result.summary
        run.external_reference = result.external_reference
        run.completed_at = datetime.now(UTC)
        self._audit.record(
            actor_id=run.created_by,
            action="COMPLETE_INTEGRATION_RUN",
            target_type="INTEGRATION_RUN",
            target_id=run.id,
            project_id=run.project_id,
            after={"status": run.status, "connector_code": run.connector_code},
        )
        self._session.commit()

    def _fail_integration_run(self, run: IntegrationRun, error_code: str, message: str) -> None:
        run.status = "FAILED"
        run.error_code = error_code
        run.error_message = message
        run.completed_at = datetime.now(UTC)
        self._audit.record(
            actor_id=run.created_by,
            action="FAIL_INTEGRATION_RUN",
            target_type="INTEGRATION_RUN",
            target_id=run.id,
            project_id=run.project_id,
            after={"status": run.status, "error_code": error_code},
        )
        self._session.commit()

    # Responses and reusable checks
    def _knowledge_response(
        self, entry: KnowledgeEntry, version: KnowledgeVersion
    ) -> KnowledgeResponse:
        source_version = (
            self._documents.get_version(version.source_document_version_id)
            if version.source_document_version_id is not None
            else None
        )
        return KnowledgeResponse(
            entry_id=entry.id,
            version_id=version.id,
            version_no=version.version_no,
            knowledge_type=entry.knowledge_type,
            title=entry.title,
            authority=entry.authority,
            source_reference=entry.source_reference,
            status=version.status,
            content=version.content,
            issued_on=version.issued_on,
            effective_on=version.effective_on,
            citation_note=version.citation_note,
            source_document_version_id=version.source_document_version_id,
            source_parse_status=None if source_version is None else source_version.parse_status,
            source_cleaning_summary=None
            if source_version is None
            else source_version.cleaning_summary,
            published_at=version.published_at,
            created_at=version.created_at,
        )

    @staticmethod
    def _manual_knowledge_evidence(
        entry: KnowledgeEntry, version: KnowledgeVersion, actor_id: UUID
    ) -> Evidence:
        content = version.content.strip()
        return Evidence(
            id=uuid4(),
            source_type="USER_CONFIRMATION",
            document_version_id=None,
            document_node_id=None,
            page_number=None,
            quoted_text=content[:1_000],
            content_hash=hashlib.sha256(content.encode("utf-8")).hexdigest(),
            bbox=None,
            source_reference={
                "knowledge_entry_id": str(entry.id),
                "knowledge_version_id": str(version.id),
                "source_reference": entry.source_reference,
            },
            confidence=None,
            created_at=datetime.now(UTC),
            created_by=actor_id,
        )

    def _analysis_response(self, analysis: CompetitiveAnalysis) -> CompetitiveAnalysisResponse:
        return CompetitiveAnalysisResponse(
            id=analysis.id,
            project_id=analysis.project_id,
            requirement_id=analysis.requirement_id,
            status=analysis.status,
            method=analysis.method,
            summary=analysis.summary,
            evidence_ids=self._analysis_evidence_ids(analysis.id),
            findings=[
                self._finding_response(item) for item in self._repo.list_findings(analysis.id)
            ],
            created_at=analysis.created_at,
            completed_at=analysis.completed_at,
        )

    def _finding_response(self, finding: CompetitiveFinding) -> CompetitiveFindingResponse:
        return CompetitiveFindingResponse(
            id=finding.id,
            category=finding.category,
            title=finding.title,
            description=finding.description,
            confidence=finding.confidence,
            status=finding.status,
            resolution=finding.resolution,
            evidence_ids=self._finding_evidence_ids(finding.id),
            knowledge_version_ids=self._finding_knowledge_ids(finding.id),
            reviewed_by=finding.reviewed_by,
            reviewed_at=finding.reviewed_at,
        )

    def _challenge_response(self, draft: ChallengeDraft) -> ChallengeDraftResponse:
        return ChallengeDraftResponse(
            id=draft.id,
            project_id=draft.project_id,
            title=draft.title,
            subject=draft.subject,
            fact_statement=draft.fact_statement,
            requested_action=draft.requested_action,
            status=draft.status,
            review_note=draft.review_note,
            reviewed_by=draft.reviewed_by,
            reviewed_at=draft.reviewed_at,
            evidence_ids=self._challenge_evidence_ids(draft.id),
            has_docx=draft.docx_object_key is not None,
            has_pdf=draft.pdf_object_key is not None,
            created_at=draft.created_at,
            updated_at=draft.updated_at,
        )

    @staticmethod
    def _quote_response(item: QuoteScenario) -> QuoteScenarioResponse:
        return QuoteScenarioResponse(
            id=item.id,
            project_id=item.project_id,
            parent_scenario_id=item.parent_scenario_id,
            name=item.name,
            version_no=item.version_no,
            status=item.status,
            cost_excluding_tax=item.cost_excluding_tax,
            tax_rate=item.tax_rate,
            target_margin_rate=item.target_margin_rate,
            risk_adjustment=item.risk_adjustment,
            expected_score=item.expected_score,
            assumptions=item.assumptions,
            calculations=item.calculations,
            locked_at=item.locked_at,
            created_at=item.created_at,
        )

    @staticmethod
    def _comment_response(item: ProjectComment) -> ProjectCommentResponse:
        return ProjectCommentResponse(
            id=item.id,
            project_id=item.project_id,
            content=item.content,
            target_type=item.target_type,
            target_id=item.target_id,
            created_by=item.created_by,
            created_at=item.created_at,
        )

    @staticmethod
    def _work_item_response(item: WorkItem) -> WorkItemResponse:
        return WorkItemResponse(
            id=item.id,
            project_id=item.project_id,
            title=item.title,
            description=item.description,
            status=item.status,
            assignee_id=item.assignee_id,
            due_at=item.due_at,
            target_type=item.target_type,
            target_id=item.target_id,
            closing_note=item.closing_note,
            created_by=item.created_by,
            created_at=item.created_at,
            updated_at=item.updated_at,
        )

    @staticmethod
    def _notification_response(item: Notification) -> NotificationResponse:
        return NotificationResponse(
            id=item.id,
            project_id=item.project_id,
            notification_type=item.notification_type,
            payload=item.payload,
            read_at=item.read_at,
            created_at=item.created_at,
        )

    @staticmethod
    def _market_check_response(item: MarketCheck) -> MarketCheckResponse:
        return MarketCheckResponse(
            id=item.id,
            project_id=item.project_id,
            requirement_id=item.requirement_id,
            evidence_id=item.evidence_id,
            parameter=item.parameter,
            source_name=item.source_name,
            source_reference=item.source_reference,
            excerpt=item.excerpt,
            conclusion=item.conclusion,
            note=item.note,
            created_at=item.created_at,
            created_by=item.created_by,
        )

    def _graph_response(self, project_id: UUID) -> ProjectGraphResponse:
        return ProjectGraphResponse(
            nodes=[
                GraphNodeResponse(
                    id=item.id,
                    entity_type=item.entity_type,
                    source_object_id=item.source_object_id,
                    label=item.label,
                    attributes=item.attributes,
                    source_evidence_id=item.source_evidence_id,
                )
                for item in self._repo.list_graph_nodes(project_id)
            ],
            edges=[
                GraphEdgeResponse(
                    id=item.id,
                    from_node_id=item.from_node_id,
                    to_node_id=item.to_node_id,
                    relation_type=item.relation_type,
                    source_evidence_id=item.source_evidence_id,
                )
                for item in self._repo.list_graph_edges(project_id)
            ],
        )

    def _agent_run_response(self, run: AgentRun) -> AgentRunResponse:
        return AgentRunResponse(
            id=run.id,
            project_id=run.project_id,
            source_document_version_id=run.source_document_version_id,
            workflow=run.workflow,
            status=run.status,
            goal=run.goal,
            input_hash=run.input_hash,
            thread_id=run.thread_id,
            checkpoint_version=run.checkpoint_version,
            requires_human_review=run.requires_human_review,
            result=run.result,
            error_code=run.error_code,
            error_message=run.error_message,
            created_at=run.created_at,
            started_at=run.started_at,
            completed_at=run.completed_at,
            evidence_ids=list(
                self._session.scalars(
                    select(AgentRunEvidence.evidence_id).where(
                        AgentRunEvidence.agent_run_id == run.id
                    )
                )
            ),
            steps=[],
        )

    def _connector_response(self, item) -> ConnectorResponse:
        return ConnectorResponse(
            code=item.code,
            name=item.name,
            capabilities=item.capabilities,
            is_enabled=item.is_enabled,
            is_configured=self._connector_is_configured(item.code),
        )

    @staticmethod
    def _integration_run_response(run: IntegrationRun) -> IntegrationRunResponse:
        return IntegrationRunResponse(
            id=run.id,
            project_id=run.project_id,
            connector_code=run.connector_code,
            operation=run.operation,
            status=run.status,
            result_summary=run.result_summary,
            external_reference=run.external_reference,
            error_code=run.error_code,
            error_message=run.error_message,
            created_at=run.created_at,
            started_at=run.started_at,
            completed_at=run.completed_at,
        )

    def _keyword_findings(
        self,
        analysis_id: UUID,
        texts: list[str],
        evidence_ids: list[UUID],
        knowledge_ids: list[UUID],
        now: datetime,
    ) -> list[object]:
        all_text = "\n".join(texts)
        findings: list[object] = []
        for category, keywords in _FINDING_CATEGORIES.items():
            matched = [keyword for keyword in keywords if keyword in all_text]
            if not matched:
                continue
            finding = CompetitiveFinding(
                id=uuid4(),
                analysis_id=analysis_id,
                category=category,
                title=_FINDING_TITLES[category],
                description=f"检测到可能需要复核的关键词：{'、'.join(matched)}。仅为风险提示，不构成法律结论。",
                confidence=Decimal("0.6000"),
                status="PENDING",
                resolution=None,
                reviewed_by=None,
                reviewed_at=None,
                created_at=now,
            )
            findings.extend([finding])
            findings.extend(
                [
                    CompetitiveFindingEvidence(finding_id=finding.id, evidence_id=item)
                    for item in evidence_ids
                ]
            )
            findings.extend(
                [
                    CompetitiveFindingKnowledge(finding_id=finding.id, knowledge_version_id=item)
                    for item in knowledge_ids
                ]
            )
        return findings

    def _llm_findings(
        self,
        analysis_id: UUID,
        texts: list[str],
        evidence_ids: list[UUID],
        knowledge_ids: list[UUID],
        now: datetime,
    ) -> list[object]:
        """使用 LLM 生成竞争性条款候选发现，所有状态为 PENDING。"""
        if self._llm is None:
            raise DomainError(
                "AI_SERVICE_UNAVAILABLE",
                "AI 服务未配置，无法使用 LLM 分析模式",
                503,
            )

        import json

        from app.schemas.advanced import CompetitiveFindingDraft

        evidence_contexts = [
            {"index": i, "quoted_text": text or "(无引用文本)"}
            for i, text in enumerate(texts)
        ]
        payload = json.dumps(
            {
                "evidence_contexts": evidence_contexts,
                "categories": list(_FINDING_CATEGORIES.keys()),
                "category_titles": _FINDING_TITLES,
            },
            ensure_ascii=False,
        )

        system = (
            "你是一个专业的招投标法律合规分析师。"
            "基于提供的证据片段，识别可能存在竞争性条款限制的风险。"
            f"风险类别包括：{', '.join(_FINDING_CATEGORIES.keys())}。"
            "只输出你确实有依据的发现，不要虚构。"
            "描述中必须引用具体的证据索引和原文内容。"
            "所有发现状态必须为 PENDING，不得输出 CONFIRMED、RESOLVED 等终态。"
            "以结构化 JSON 列表形式输出 findings 数组。"
        )
        user = (
            "请分析以下证据片段，识别竞争性条款风险：\n\n"
            f"{payload}\n\n"
            "输出格式：\n"
            "{\n"
            '  "findings": [\n'
            '    {"category": "类别", "title": "标题", '
            '"description": "描述（需引用证据索引和原文）", '
            '"confidence": 0.0-1.0, "evidence_indices": [证据索引列表]},\n'
            "    ...\n"
            "  ]\n"
            "}"
        )

        try:
            raw = self._llm.generate_raw(
                [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ]
            )
            parsed = json.loads(raw)
            drafts: list[CompetitiveFindingDraft] = []
            for item in parsed.get("findings", []):
                try:
                    drafts.append(CompetitiveFindingDraft(**item))
                except Exception:
                    continue
        except Exception as exc:
            raise DomainError(
                "AI_SERVICE_FAILED",
                f"LLM 分析失败：{exc}",
                503,
            ) from exc

        findings: list[object] = []
        for draft in drafts:
            # 安全地取 evidence_ids，防止越界
            linked_evidence = [
                evidence_ids[i] for i in draft.evidence_indices if 0 <= i < len(evidence_ids)
            ]
            if not linked_evidence:
                linked_evidence = list(evidence_ids)
            finding = CompetitiveFinding(
                id=uuid4(),
                analysis_id=analysis_id,
                category=draft.category,
                title=draft.title,
                description=f"{draft.description}（LLM 置信度：{draft.confidence:.2f}）",
                confidence=Decimal(str(draft.confidence)),
                status="PENDING",
                resolution=None,
                reviewed_by=None,
                reviewed_at=None,
                created_at=now,
            )
            findings.append(finding)
            findings.extend(
                CompetitiveFindingEvidence(finding_id=finding.id, evidence_id=eid)
                for eid in linked_evidence
            )
            findings.extend(
                CompetitiveFindingKnowledge(finding_id=finding.id, knowledge_version_id=kid)
                for kid in knowledge_ids
            )
        return findings

    def _deduplicate_findings(
        self,
        keyword_findings: list[object],
        llm_findings: list[object],
    ) -> list[object]:
        """合并关键词和 LLM 发现，按 category 去重，保留 LLM 高置信度结果。"""
        by_category: dict[str, object] = {}
        for f in keyword_findings:
            cat = getattr(f, "category", None)
            if cat:
                by_category[cat] = f
        for f in llm_findings:
            cat = getattr(f, "category", None)
            if cat is None:
                continue
            existing = by_category.get(cat)
            if existing is None:
                by_category[cat] = f
            else:
                # 同类别：取置信度更高的
                existing_conf = getattr(existing, "confidence", Decimal("0")) or Decimal("0")
                new_conf = getattr(f, "confidence", Decimal("0")) or Decimal("0")
                if new_conf > existing_conf:
                    by_category[cat] = f
        return list(by_category.values())

    def _validated_project_evidence_ids(
        self, project_id: UUID, evidence_ids: list[UUID]
    ) -> list[UUID]:
        unique_ids = list(dict.fromkeys(evidence_ids))
        for evidence_id in unique_ids:
            evidence = self._evidences.get(evidence_id)
            if evidence is None or evidence.document_version_id is None:
                raise DomainError("RESOURCE_NOT_FOUND", "Evidence 不存在", 404)
            version = self._documents.get_version(evidence.document_version_id)
            document = self._documents.get_document(version.document_id) if version else None
            if (
                document is None
                or document.document_type != "TENDER"
                or document.project_id != project_id
            ):
                raise DomainError("RESOURCE_NOT_FOUND", "Evidence 不属于当前项目", 404)
        return unique_ids

    def _published_knowledge_ids(self, version_ids: list[UUID]) -> list[UUID]:
        unique_ids = list(dict.fromkeys(version_ids))
        for version_id in unique_ids:
            version = self._repo.get_knowledge_version(version_id)
            if version is None or version.status != "PUBLISHED":
                raise DomainError("RESOURCE_NOT_FOUND", "已发布知识版本不存在", 404)
        return unique_ids

    def _analysis_evidence_ids(self, analysis_id: UUID) -> list[UUID]:
        return list(
            self._session.scalars(
                select(CompetitiveAnalysisEvidence.evidence_id).where(
                    CompetitiveAnalysisEvidence.analysis_id == analysis_id
                )
            )
        )

    def _finding_evidence_ids(self, finding_id: UUID) -> list[UUID]:
        return list(
            self._session.scalars(
                select(CompetitiveFindingEvidence.evidence_id).where(
                    CompetitiveFindingEvidence.finding_id == finding_id
                )
            )
        )

    def _finding_knowledge_ids(self, finding_id: UUID) -> list[UUID]:
        return list(
            self._session.scalars(
                select(CompetitiveFindingKnowledge.knowledge_version_id).where(
                    CompetitiveFindingKnowledge.finding_id == finding_id
                )
            )
        )

    def _challenge_evidence_ids(self, draft_id: UUID) -> list[UUID]:
        return list(
            self._session.scalars(
                select(ChallengeDraftEvidence.evidence_id).where(
                    ChallengeDraftEvidence.challenge_draft_id == draft_id
                )
            )
        )

    def _require_project_writer(
        self, project_id: UUID, actor_id: UUID, role_codes: set[str]
    ) -> None:
        project = self._projects.get_visible(project_id, actor_id, role_codes)
        self._projects.require_writable(project)
        if not _WRITER_ROLES.intersection(role_codes):
            raise DomainError("PERMISSION_DENIED", "无权执行项目高级分析操作", 403)

    @staticmethod
    def _require_system_admin(role_codes: set[str]) -> None:
        if SYSTEM_ADMIN not in role_codes:
            raise DomainError("PERMISSION_DENIED", "仅系统管理员可执行该操作", 403)

    @staticmethod
    def _can_manage_knowledge(role_codes: set[str]) -> bool:
        return bool({SYSTEM_ADMIN, LEGAL_COMPLIANCE}.intersection(role_codes))

    def _require_knowledge_manager(self, role_codes: set[str]) -> None:
        if not self._can_manage_knowledge(role_codes):
            raise DomainError("PERMISSION_DENIED", "无权维护法规/案例知识库", 403)

    @staticmethod
    def _require_legal_reviewer(role_codes: set[str]) -> None:
        if not {SYSTEM_ADMIN, LEGAL_COMPLIANCE}.intersection(role_codes):
            raise DomainError("PERMISSION_DENIED", "仅法务/合规或管理员可审核", 403)

    def _is_project_assignee(self, project_id: UUID, user_id: UUID) -> bool:
        user = self._session.get(User, user_id)
        if user is None or user.status != "ACTIVE":
            return False
        return (
            self._session.scalar(
                select(ProjectMember.project_id).where(
                    ProjectMember.project_id == project_id,
                    ProjectMember.user_id == user_id,
                )
            )
            is not None
        )

    def _validate_target(
        self, project_id: UUID, target_type: str | None, target_id: UUID | None
    ) -> None:
        if target_type is None and target_id is None:
            return
        if not target_type or target_id is None:
            raise DomainError("VALIDATION_ERROR", "关联对象类型和 ID 必须同时提供", 422)
        if target_type == "EVIDENCE":
            self._validated_project_evidence_ids(project_id, [target_id])
            return
        if target_type == "REQUIREMENT":
            requirement = self._requirements.get(target_id)
            if requirement is not None and requirement.project_id == project_id:
                return
        if target_type == "RISK":
            risk = self._risks.get(target_id)
            if risk is not None and risk.project_id == project_id:
                return
        model = _TARGET_MODELS.get(target_type)
        target = self._session.get(model, target_id) if model else None
        if target is not None and target.project_id == project_id:
            return
        raise DomainError("RESOURCE_NOT_FOUND", "关联对象不存在或不属于当前项目", 404)

    def _upsert_graph_node(
        self,
        project_id: UUID,
        entity_type: str,
        source_object_id: str,
        label: str,
        attributes: dict[str, object],
        evidence_id: UUID | None,
        now: datetime,
    ) -> GraphNode:
        node = self._repo.get_graph_node(project_id, entity_type, source_object_id)
        if node is None:
            node = GraphNode(
                id=uuid4(),
                project_id=project_id,
                entity_type=entity_type,
                source_object_id=source_object_id,
                label=label,
                attributes=attributes,
                source_evidence_id=evidence_id,
                created_at=now,
                updated_at=now,
            )
            self._repo.add(node)
            self._session.flush()
        else:
            node.label = label
            node.attributes = attributes
            node.source_evidence_id = evidence_id
            node.updated_at = now
        return node

    def _upsert_graph_edge(
        self,
        project_id: UUID,
        from_node_id: UUID,
        to_node_id: UUID,
        relation_type: str,
        evidence_id: UUID | None,
        actor_id: UUID,
        now: datetime,
    ) -> None:
        edge = self._repo.find_graph_edge(project_id, from_node_id, to_node_id, relation_type)
        if edge is None:
            self._repo.add(
                GraphEdge(
                    id=uuid4(),
                    project_id=project_id,
                    from_node_id=from_node_id,
                    to_node_id=to_node_id,
                    relation_type=relation_type,
                    source_evidence_id=evidence_id,
                    created_at=now,
                    created_by=actor_id,
                )
            )
        else:
            edge.source_evidence_id = evidence_id

    @staticmethod
    def _quote_calculations(payload: QuoteScenarioCreate) -> dict[str, str]:
        net_quote = (payload.cost_excluding_tax + payload.risk_adjustment) / (
            Decimal("1") - payload.target_margin_rate
        )
        tax = net_quote * payload.tax_rate
        gross_profit = net_quote - payload.cost_excluding_tax
        adjusted_profit = gross_profit - payload.risk_adjustment

        def round_money(value: Decimal) -> str:
            return str(value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))

        def round_rate(value: Decimal) -> str:
            return str(value.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP))

        scenarios: dict[str, str] = {}
        for label, offset in (
            ("low", Decimal("-0.05")),
            ("baseline", Decimal("0")),
            ("high", Decimal("0.05")),
        ):
            rate = min(max(payload.target_margin_rate + offset, Decimal("0")), Decimal("0.9900"))
            scenarios[label] = round_money(
                (payload.cost_excluding_tax + payload.risk_adjustment) / (Decimal("1") - rate)
            )
        return {
            "quote_excluding_tax": round_money(net_quote),
            "tax_amount": round_money(tax),
            "quote_including_tax": round_money(net_quote + tax),
            "gross_profit": round_money(gross_profit),
            "adjusted_profit": round_money(adjusted_profit),
            "gross_margin_rate": round_rate(gross_profit / net_quote),
            "sensitivity_quotes": scenarios,
        }

    @staticmethod
    def _agent_result(
        workflow: str, requirements, risks, evidence_ids: list[UUID]
    ) -> dict[str, object]:
        pending = [risk for risk in risks if risk.status in {"PENDING", "CONFIRMED"}]
        critical = [risk for risk in pending if risk.severity == "CRITICAL"]
        checklist = [
            {"action": "复核已确认 Requirement", "count": len(requirements)},
            {"action": "处理待办风险提示", "count": len(pending)},
            {"action": "核对引用 Evidence", "count": len(evidence_ids)},
        ]
        if workflow == "MARKET_REVIEW":
            checklist.append(
                {"action": "对关键参数建立人工市场核查记录", "count": len(requirements)}
            )
        if workflow == "COMPLIANCE_REVIEW":
            checklist.append({"action": "对竞争性条款发现进行法务复核", "count": len(critical)})
        return {
            "execution_mode": "DETERMINISTIC_EVIDENCE_WORKFLOW",
            "requires_manual_adoption": True,
            "warning": "结果为待人工采纳的工作建议，不会自动修改任何业务对象或外部系统。",
            "steps": checklist,
            "evidence_ids": [str(item) for item in evidence_ids],
        }

    def _connector_is_configured(self, code: str) -> bool:
        return bool(
            {
                "ERP": self._settings.erp_integration_base_url,
                "CRM": self._settings.crm_integration_base_url,
                "PUBLIC_RESOURCE": self._settings.public_resource_integration_base_url,
            }.get(code)
        )

    @staticmethod
    def _hash(value: object) -> str:
        serialized = json.dumps(
            value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str
        )
        return hashlib.sha256(serialized.encode("utf-8")).hexdigest()

    @staticmethod
    def _render_challenge_docx(draft: ChallengeDraft, evidence_ids: list[UUID]) -> bytes:
        document = DocxDocument()
        document.styles["Normal"].font.name = "Microsoft YaHei"
        document.styles["Normal"].font.size = Pt(10.5)
        document.add_heading("质疑函草稿（仅供人工法务审核）", 0)
        document.add_paragraph("系统不会发送、提交或代表任何主体作出法律结论。")
        for heading, value in (
            ("标题", draft.title),
            ("事项", draft.subject),
            ("事实说明", draft.fact_statement),
            ("请求事项", draft.requested_action),
        ):
            document.add_heading(heading, 1)
            document.add_paragraph(value)
        document.add_heading("Evidence 索引", 1)
        for evidence_id in evidence_ids:
            document.add_paragraph(str(evidence_id), style="List Bullet")
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _render_challenge_pdf(draft: ChallengeDraft, evidence_ids: list[UUID]) -> bytes:
        buffer = BytesIO()
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        pdf = canvas.Canvas(buffer, pagesize=A4)
        pdf.setFont("STSong-Light", 14)
        y = 800
        for line in [
            "质疑函草稿（仅供人工法务审核）",
            "系统不会发送、提交或代表任何主体作出法律结论。",
            f"标题：{draft.title}",
            f"事项：{draft.subject}",
            f"事实说明：{draft.fact_statement}",
            f"请求事项：{draft.requested_action}",
            "Evidence 索引：",
            *[str(item) for item in evidence_ids],
        ]:
            for segment in AdvancedService._pdf_lines(line, 42):
                if y < 48:
                    pdf.showPage()
                    pdf.setFont("STSong-Light", 10)
                    y = 800
                pdf.drawString(48, y, segment)
                y -= 18
        pdf.save()
        return buffer.getvalue()

    @staticmethod
    def _pdf_lines(value: str, width: int) -> list[str]:
        return [value[index : index + width] for index in range(0, len(value), width)] or [""]

    def _compensate_object(self, object_key: str) -> None:
        try:
            self._storage.delete_object(object_key)
        except ObjectStorageUnavailable:
            pass
