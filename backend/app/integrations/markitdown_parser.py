"""MarkItDown 本地文档解析器 - 作为 MinerU 的 fallback 方案。

支持格式：DOCX, PPTX, XLSX, PDF 等
"""

import io
import logging
import re
from pathlib import Path

from markitdown import MarkItDown

from app.integrations.mineru import MinerUParseResult, ParsedNode, ParserUnavailable

logger = logging.getLogger(__name__)


class MarkItdownClient:
    """本地文档解析器，使用 markitdown 库直接解析，无需外部 API。"""

    def __init__(self) -> None:
        self._markitdown = MarkItDown()

    def parse(self, source_path: Path, source_mime_type: str | None = None) -> MinerUParseResult:
        """使用 MarkItDown 解析文档。

        Args:
            source_path: 文档路径
            source_mime_type: MIME 类型（可选，用于提示格式）

        Returns:
            MinerUParseResult 格式的结果

        Raises:
            ParserUnavailable: 解析失败时
        """
        ext = source_path.suffix.lstrip(".").lower()
        try:
            with source_path.open("rb") as f:
                content = f.read()
            result = self._markitdown.convert(
                io.BytesIO(content),
                file_extension=ext or None,
                keep_data_uris=False,
            )
            text = result.text_content or ""
        except Exception as exc:
            if ext == "docx":
                logger.warning(
                    "MarkItDown DOCX conversion failed; using python-docx fallback: %s", exc
                )
                text = self._parse_docx(source_path)
            else:
                raise ParserUnavailable(f"MarkItDown 解析失败: {exc}") from exc

        nodes = self._text_to_nodes(text)
        if not nodes:
            raise ParserUnavailable("MarkItDown 返回空内容")

        return MinerUParseResult(
            nodes=tuple(nodes),
            raw_output=text.encode("utf-8"),
            raw_content_type="text/plain",
        )

    @staticmethod
    def _parse_docx(source_path: Path) -> str:
        """Fallback for DOCX when MarkItDown optional converters are unavailable."""
        try:
            from docx import Document

            document = Document(source_path)
            parts = [paragraph.text.strip() for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            text = "\n".join(part for part in parts if part)
        except Exception as exc:
            raise ParserUnavailable(f"DOCX fallback 解析失败: {exc}") from exc
        if not text:
            raise ParserUnavailable("DOCX fallback 返回空内容")
        return text

    def _text_to_nodes(self, text: str) -> list[ParsedNode]:
        """将文本转换为 ParsedNode 列表。

        简单的标题识别（# 开头）和段落分离。
        复杂布局信息会丢失，但保留了核心内容。
        """
        nodes: list[ParsedNode] = []
        section_path: list[str] = []
        order_no = 0

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue

            order_no += 1

            # 识别标题（# 开头或全大写短行）
            if line.startswith("#"):
                # Markdown 标题
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()
                if title:
                    section_path = [title]
                    nodes.append(
                        ParsedNode(
                            node_type="SECTION",
                            content=title,
                            page_number=None,
                            section_path=title if level == 1 else " / ".join(section_path),
                            metadata={"parser": "markitdown", "heading_level": level},
                        )
                    )
            elif self._is_likely_heading(line):
                # 识别为标题（非 Markdown 格式）
                section_path.append(line)
                nodes.append(
                    ParsedNode(
                        node_type="SECTION",
                        content=line,
                        page_number=None,
                        section_path=" / ".join(section_path[-3:]),
                        metadata={"parser": "markitdown"},
                    )
                )
            else:
                # 普通段落
                nodes.append(
                    ParsedNode(
                        node_type="PARAGRAPH",
                        content=line,
                        page_number=None,
                        section_path=" / ".join(section_path) if section_path else None,
                        metadata={"parser": "markitdown"},
                    )
                )

        return nodes

    @staticmethod
    def _is_likely_heading(line: str) -> bool:
        r"""判断一行文本是否可能是标题。

        参考 WeKnora patterns.go / profiler.go 的纯行级正则策略：

        1. 行首锚定的强信号（任一命中即 True）：
           - 中文章节：第X(章|节|部分|篇)（不含"条"，避免误判"第一条 招标范围"是正文）
           - 数字编号：\\d+. 或 一、 / （一）
           - 段落标记：（一）（二）

        2. 行首关键词 + 行短 + 不以句号结尾（防止"本项目采购货物..."误判）：
           - 仅识别 2-30 字符行
           - 行首必须是 ["项目概况", "项目预算", "工程概况", ...]
             这类典型**名词性**标题词，不能含动词/连接词

        3. 排除条件（中文）：不以句末标点结尾（。！？），避免把含项目/采购的短正文误判
        """
        stripped = line.strip()
        if not stripped or len(stripped) > 50:
            return False

        # 强信号 1：中文章节（第X章/节/部分/篇，不含"条"）
        if re.match(r"^第[一二三四五六七八九十百千零〇\d]+(章|节|部分|篇)\b", stripped):
            return True

        # 强信号 2：数字编号标题（如 "1. xxx" / "一、 xxx" / "（一）xxx"）
        if re.match(r"^\d+\.\s+\S", stripped):
            return True
        if re.match(r"^[一二三四五六七八九十]+[、．]\s*\S", stripped):
            return True
        if re.match(r"^（[一二三四五六七八九十\d]+）\s*\S", stripped):
            return True

        # 弱信号：行首必须是"典型名词性标题"，且不含动词/连接词
        if stripped.endswith(("。", "！", "？", ".", "!", "?")):
            return False

        # 行首严格匹配（不是"包含"），防止"本项目..."误判
        noun_heading_prefixes = (
            "项目概况", "项目预算", "项目内容", "项目范围",
            "工程概况", "工程内容", "工程范围",
            "招标范围", "招标内容", "招标方式",
            "采购方式", "采购内容", "采购范围", "采购需求",
            "投标须知", "投标要求", "投标人须知",
            "合同条款", "合同主要条款", "合同格式",
            "服务内容", "服务要求", "服务范围",
            "评标办法", "评审办法",
            "投标人资格", "投标人要求", "资格审查",
            "技术要求", "技术规格", "技术标准",
        )
        # 长度 ≤ 30 字符 + 行首是已知名词性标题词前缀
        if len(stripped) <= 30 and any(stripped.startswith(p) for p in noun_heading_prefixes):
            return True

        return False


def create_markitdown_client() -> MarkItdownClient:
    """工厂函数：创建 MarkItDown 客户端。"""
    return MarkItdownClient()
